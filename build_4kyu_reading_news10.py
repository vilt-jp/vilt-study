# -*- coding: utf-8 -*-
"""4級読解 時事対策10問。5分野×各短文1本×2問=10問。
構成: ①やさしい越語本文(5本) ②読解10問+解答 ③(改ページ)日本語訳 ④4級単語・文法。
フォントMeiryo UI(latin/ea/cs)。練習用(実在の特定記事ではない・傾向ベース)。"""
from docx import Document
from docx.shared import Pt, RGBColor
from docx.oxml.ns import qn
from docx.enum.text import WD_ALIGN_PARAGRAPH

FONT="Meiryo UI"
def setf(run,size=None,bold=None,color=None):
    run.font.name=FONT
    rPr=run._element.get_or_add_rPr()
    rF=rPr.find(qn('w:rFonts'))
    if rF is None:
        rF=rPr.makeelement(qn('w:rFonts'),{}); rPr.insert(0,rF)
    for a in ('w:ascii','w:hAnsi','w:eastAsia','w:cs'): rF.set(qn(a),FONT)
    if size is not None: run.font.size=Pt(size)
    if bold is not None: run.font.bold=bold
    if color is not None: run.font.color.rgb=RGBColor(*color)
def para(doc,text="",size=11,bold=False,color=None,align=None,after=6):
    p=doc.add_paragraph()
    if align is not None: p.alignment=align
    p.paragraph_format.space_after=Pt(after)
    setf(p.add_run(text),size=size,bold=bold,color=color); return p
def head(doc,text,size=14):
    return para(doc,text,size=size,bold=True,color=(0x1F,0x49,0x7D),after=8)

doc=Document()
st=doc.styles['Normal']; st.font.name=FONT; st.element.rPr.rFonts.set(qn('w:eastAsia'),FONT)

para(doc,"実用ベトナム語検定 4級 読解対策 ― 時事の話題10問",size=15,bold=True,
     color=(0x1F,0x49,0x7D),align=WD_ALIGN_PARAGRAPH.CENTER,after=4)
para(doc,"5分野×短文5本×各2問＝10問　／　練習用（実在の特定記事ではなく、最近の傾向をもとにした文）",
     size=9.5,color=(0x66,0x66,0x66),align=WD_ALIGN_PARAGRAPH.CENTER,after=10)

# 5本の短文（分野, 越語本文, 日本語訳, [(問,選択肢4,正解index)])
passages=[
("【分野1 経済・社会】 Thanh toán không dùng tiền mặt（キャッシュレス決済）",
 "Hiện nay, ngày càng nhiều người Việt Nam thanh toán không dùng tiền mặt. Khi mua đồ ở siêu thị, quán cà phê hay chợ, nhiều người chỉ cần dùng điện thoại để quét mã QR. Cách này nhanh và tiện. Tuy nhiên, một số người lớn tuổi vẫn thích dùng tiền mặt vì đã quen.",
 "今、ますます多くのベトナム人が現金を使わずに支払いをしている。スーパーや喫茶店、市場で買い物をするとき、多くの人は電話を使ってQRコードを読み取るだけでよい。この方法は速くて便利だ。ただし、一部の年配の人は、慣れているので今も現金を使うのが好きだ。",
 [("Câu 1. Nhiều người Việt Nam bây giờ thanh toán bằng cách nào?",
   ["A. Trả bằng tiền mặt","B. Quét mã QR bằng điện thoại","C. Gửi thư","D. Gọi điện thoại"],1),
  ("Câu 2. Tại sao một số người lớn tuổi vẫn thích tiền mặt?",
   ["A. Vì rẻ hơn","B. Vì nhanh hơn","C. Vì đã quen","D. Vì an toàn hơn"],2)]),
("【分野2 くらし・働き方】 Làm việc ở nhà（在宅勤務）",
 "Sau dịch Covid-19, nhiều công ty cho nhân viên làm việc ở nhà. Nhân viên không phải đi lại nên tiết kiệm được thời gian. Nhưng làm việc ở nhà cũng khó vì đôi khi khó tập trung. Vì vậy, nhiều người chọn cách vừa làm ở nhà vừa lên công ty.",
 "新型コロナの流行のあと、多くの会社が従業員を在宅で働かせている。従業員は通勤しなくてよいので、時間を節約できる。しかし、在宅勤務は、時々集中しにくいので難しさもある。そのため、多くの人は、家でも働き会社にも行く、という方法を選んでいる。",
 [("Câu 3. Làm việc ở nhà có lợi gì?",
   ["A. Tiết kiệm thời gian đi lại","B. Được nhiều tiền hơn","C. Gặp nhiều bạn hơn","D. Ăn ngon hơn"],0),
  ("Câu 4. Khó khăn khi làm việc ở nhà là gì?",
   ["A. Không có máy tính","B. Đôi khi khó tập trung","C. Nhà quá xa","D. Không có điện"],1)]),
("【分野3 文化・行事・観光】 Du lịch（観光）",
 "Năm nay, rất nhiều khách du lịch nước ngoài đến Việt Nam. Các thành phố như Đà Nẵng, Hà Nội và Hội An rất đông khách. Khách thích ăn món ăn Việt Nam và đi thăm các nơi đẹp. Ngành du lịch vì thế phát triển trở lại.",
 "今年は、とても多くの外国人観光客がベトナムに来ている。ダナン、ハノイ、ホイアンといった都市は観光客でとても混んでいる。観光客はベトナム料理を食べたり、美しい場所を訪れたりするのが好きだ。観光業はそのため、再び発展している。",
 [("Câu 5. Khách du lịch thích làm gì ở Việt Nam?",
   ["A. Làm việc","B. Học tiếng Việt","C. Ăn món Việt và đi thăm nơi đẹp","D. Mua nhà"],2),
  ("Câu 6. Thành phố nào đông khách du lịch?",
   ["A. Đà Nẵng, Hà Nội, Hội An","B. Tokyo","C. Seoul","D. Paris"],0)]),
("【分野4 会社・人物】 Xe điện（電気自動車）",
 "VinFast là một công ty Việt Nam sản xuất xe ô tô chạy bằng điện. Xe điện không dùng xăng nên không thải khói, tốt cho môi trường. Hiện nay, nhiều người Việt bắt đầu mua xe điện. Công ty cũng bán xe sang nước ngoài.",
 "ビンファストは、電気で走る自動車を生産するベトナムの会社だ。電気自動車はガソリンを使わないので煙を出さず、環境に良い。今、多くのベトナム人が電気自動車を買い始めている。会社は外国にも車を売っている。",
 [("Câu 7. Vì sao xe điện tốt cho môi trường?",
   ["A. Vì chạy nhanh","B. Vì không thải khói","C. Vì rất rẻ","D. Vì rất to"],1),
  ("Câu 8. VinFast là công ty của nước nào?",
   ["A. Nhật Bản","B. Mỹ","C. Hàn Quốc","D. Việt Nam"],3)]),
("【分野5 サービス・客】 Giao đồ ăn（料理の宅配）",
 "Bây giờ, nhiều người Việt đặt đồ ăn qua điện thoại. Chỉ cần dùng ứng dụng, chọn món và chờ ở nhà. Người giao hàng sẽ mang đồ ăn đến tận nơi. Dịch vụ này rất tiện cho người bận rộn, nhưng tiền giao hàng đôi khi hơi cao.",
 "今、多くのベトナム人が電話で料理を注文している。アプリを使い、料理を選び、家で待つだけでよい。配達の人が料理をその場所まで持ってきてくれる。このサービスは忙しい人にとても便利だが、配達料が時々少し高い。",
 [("Câu 9. Người ta đặt đồ ăn bằng cách nào?",
   ["A. Dùng ứng dụng trên điện thoại","B. Đi đến quán","C. Gửi thư","D. Tự nấu"],0),
  ("Câu 10. Nhược điểm của dịch vụ này là gì?",
   ["A. Đồ ăn không ngon","B. Phải chờ rất lâu","C. Tiền giao hàng đôi khi hơi cao","D. Không có món để chọn"],2)]),
]

# ① 本文
head(doc,"Ⅰ. Bài đọc (やさしいベトナム語・5本)")
for title,vn,jp,qs in passages:
    para(doc,title,size=11.5,bold=True,after=2)
    para(doc,vn,size=11,after=8)

# ② 設問
doc.add_paragraph()
head(doc,"Ⅱ. Câu hỏi đọc hiểu (読解問題・全10問)")
para(doc,"上の本文を読み、正しい答えを1つ選びなさい。",size=10,color=(0x66,0x66,0x66),after=8)
answers=[]
for ti,(title,vn,jp,qs) in enumerate(passages,1):
    para(doc,title.split("】")[0]+"】",size=10.5,bold=True,color=(0x55,0x55,0x55),after=2)
    for q,opts,ai in qs:
        para(doc,q,size=11,bold=True,after=2)
        for j,o in enumerate(opts): para(doc,"   "+o,size=11,after=1)
        # 正解記録
        num=q.split(".")[0].replace("Câu ","")
        answers.append((num,opts[ai].split(".")[0]))
        doc.add_paragraph()
ans_line="【解答】 "+"  /  ".join("%s: %s"%(n,a) for n,a in answers)
para(doc,ans_line,size=11,bold=True,color=(0xC0,0,0),after=6)

# ③ 日本語訳
doc.add_page_break()
head(doc,"Ⅲ. 日本語訳（本文5本）")
for title,vn,jp,qs in passages:
    para(doc,title,size=11.5,bold=True,after=2)
    para(doc,jp,size=11,after=8)

# ④ 4級単語・文法
doc.add_paragraph()
head(doc,"Ⅳ. 4級の単語・文法")
para(doc,"▼ 単語",size=12,bold=True,color=(0x1F,0x49,0x7D),after=4)
vocab=[
("thanh toán","支払う"),("tiền mặt","現金"),("quét mã QR","QRコードを読み取る"),
("tiện","便利"),("người lớn tuổi","年配の人"),("quen","慣れている"),
("nhân viên","従業員"),("làm việc ở nhà","在宅で働く"),("tiết kiệm","節約する"),
("đi lại","行き来する・通勤する"),("tập trung","集中する"),("khách du lịch","観光客"),
("nước ngoài","外国"),("đông","混んでいる"),("phát triển","発展する"),
("xe điện","電気自動車"),("sản xuất","生産する"),("xăng","ガソリン"),
("thải khói","煙を出す"),("môi trường","環境"),("đặt đồ ăn","料理を注文する"),
("ứng dụng","アプリ"),("giao hàng","配達する"),("tận nơi","その場所まで"),
("bận rộn","忙しい"),("nhược điểm","欠点・短所"),
]
tbl=doc.add_table(rows=1,cols=2); tbl.style='Table Grid'
for c,t in zip(tbl.rows[0].cells,["ベトナム語","意味"]):
    setf(c.paragraphs[0].add_run(t),size=10.5,bold=True)
for vi,jp_ in vocab:
    cells=tbl.add_row().cells
    setf(cells[0].paragraphs[0].add_run(vi),size=11)
    setf(cells[1].paragraphs[0].add_run(jp_),size=11)

doc.add_paragraph()
para(doc,"▼ 文法（やさしい言い方）",size=12,bold=True,color=(0x1F,0x49,0x7D),after=4)
gram=[
("ngày càng + 形容詞","「ますます〜」。例: ngày càng nhiều người（ますます多くの人）"),
("chỉ cần + 動詞","「〜するだけでよい」。例: chỉ cần dùng ứng dụng（アプリを使うだけでよい）"),
("vì vậy / vì thế","「だから・そのため」。前の理由を受けて結果を言う。"),
("… nên …","「〜なので…」。例: không dùng xăng nên không thải khói（ガソリンを使わないので煙を出さない）"),
("vừa A vừa B","「AしながらB／AでもありBでもある」。例: vừa làm ở nhà vừa lên công ty"),
("tuy nhiên / nhưng","「しかし・ただし」。前と逆のことを続ける。"),
("bằng + 手段","「〜で（手段）」。例: thanh toán bằng tiền mặt（現金で支払う）"),
("cho + 人 + 動詞","「（人）に〜させる」。例: công ty cho nhân viên làm việc ở nhà"),
]
for pat,exp in gram:
    para(doc,"・"+pat,size=11,bold=True,after=1)
    para(doc,"    "+exp,size=10.5,after=6)

out=r"C:\Users\baru1\Desktop\ベトナム語\ベトナム語検定\4級\4級読解対策_時事10問.docx"
doc.save(out)
print("SAVED:",out)
print("本文",len(passages),"設問",sum(len(p[3]) for p in passages),"単語",len(vocab),"文法",len(gram))
print("解答:",ans_line)
