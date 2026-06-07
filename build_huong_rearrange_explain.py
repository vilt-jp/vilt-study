# -*- coding: utf-8 -*-
"""Huong先生 並び替え35問の「文法解説＋単語(漢越語つき)」をWord化。Goodnote取込用。
フォントMeiryo UI(latin/ea/cs)で越日両対応。"""
from docx import Document
from docx.shared import Pt, RGBColor
from docx.oxml.ns import qn
from docx.enum.text import WD_ALIGN_PARAGRAPH

FONT="Meiryo UI"
def setf(run,size=None,bold=None,color=None):
    run.font.name=FONT
    rPr=run._element.get_or_add_rPr()
    rF=rPr.find(qn('w:rFonts'))
    if rF is None:
        rF=rPr.makeelement(qn('w:rFonts'),{}); rPr.insert(0,rF)
    for a in ('w:ascii','w:hAnsi','w:eastAsia','w:cs'): rF.set(qn(a),FONT)
    if size is not None: run.font.size=Pt(size)
    if bold is not None: run.font.bold=bold
    if color is not None: run.font.color.rgb=RGBColor(*color)

def para(doc,text="",size=11,bold=False,color=None,align=None,after=5,before=0):
    p=doc.add_paragraph()
    if align is not None: p.alignment=align
    p.paragraph_format.space_after=Pt(after); p.paragraph_format.space_before=Pt(before)
    r=p.add_run(text); setf(r,size=size,bold=bold,color=color); return p

def head(doc,text,size=14):
    return para(doc,text,size=size,bold=True,color=(0x1F,0x49,0x7D),after=8,before=8)

doc=Document()
st=doc.styles['Normal']; st.font.name=FONT; st.element.rPr.rFonts.set(qn('w:eastAsia'),FONT)

para(doc,"Huong先生 並び替え35問 ― 文法と単語（漢越語つき）",size=16,bold=True,
     color=(0x1F,0x49,0x7D),align=WD_ALIGN_PARAGRAPH.CENTER,after=4)
para(doc,"4級対策・ベトナム語学習用",size=9.5,color=(0x66,0x66,0x66),
     align=WD_ALIGN_PARAGRAPH.CENTER,after=10)

# ===== 文法 =====
head(doc,"Ⅰ. 文法パターン")
grammar=[
("1. 〜だけれども（譲歩）：Dù / Mặc dù / Tuy … (nhưng) … vẫn / cũng",
 "逆のことが起きても「それでも〜」。後ろに vẫn（それでも）や cũng（〜も）が来る。Tuy が一番かたい言い方。",
 ["第1問 Dù gia đình có phản đối thì tôi vẫn cưới cô ấy.（家族が反対しても、それでも結婚する）",
  "第26問 Dù bố mẹ không đồng ý nhưng Lan vẫn đi du học.",
  "第27問 Mặc dù thời tiết xấu nhưng trận đấu vẫn được diễn ra. ／ 第28問 Tuy nhà xa nhưng…"]),
("2. 〜であろうとなかろうと：Dù … hay (không) … thì … cũng",
 "二つのどちらでも結果は同じ、という言い方。",
 ["第30問 …dù thích hay không thì chúng ta nên…（好きでも嫌いでも〜すべき）",
  "第31問 Dù bận rộn hay rảnh rỗi thì tôi cũng…"]),
("3. do（行為者）＋動詞 ＝ 〜によって…される",
 "英語の受け身に近い。「誰がやったか」を do の後ろに置く。",
 ["第3問 Mì ăn liền do Ando Momofuku phát minh ra.（即席麺は安藤百福によって発明された）",
  "第4問 Bản báo cáo này do chị Vy viết."]),
("4. 〜なので…（理由→結果）：(Do/Vì) … nên …",
 "前に理由、nên のあとに結果を置く。",
 ["第35問 Do lười học nên tôi phải thi lại…",
  "第34問 …bỏ bữa sáng nên bị đau dạ dày."]),
("5. 〜だけでなく…も：không những … mà còn … (nữa)",
 "二つの良い点・特徴を並べる。",
 ["第5問 Anh Kimura không những nấu được món Nhật mà còn món ăn Trung Quốc nữa.",
  "第6問 Trời không những mưa mà còn lạnh nữa.",
  "第7問 Cô ấy không những là một diễn viên mà còn là một nhà văn.（là が2回要る）"]),
("6. 例えば：ví dụ như",
 "例を挙げるときの言い方。",
 ["第8問 Hà Nội có nhiều hồ … ví dụ như Hồ Tây, Hồ Hoàn Kiếm."]),
("7. どんな〜でも：bất cứ + 名詞 + nào cũng / bất cứ điều gì / bất cứ thứ gì / bất cứ lúc nào",
 "「どれでも・何でも・いつでも」と全部を指す。",
 ["第9問 …bất cứ căn bếp nào cũng có nước mắm.（どの台所にも魚醤がある）",
  "第10問 …làm bất cứ điều gì…　第11問 …bất cứ thứ gì vào bất cứ lúc nào."]),
("8. 〜次第・〜に応じて：tùy theo / tùy thuộc vào（成句 tùy cơ ứng biến）",
 "条件によって結果が変わる、という言い方。",
 ["第13問 Tùy theo trọng lượng và kích thước, giá tiền thay đổi.",
  "第14問 Câu trả lời tùy thuộc vào cách suy nghĩ của mỗi người.",
  "第12問 Tùy cơ ứng biến.（臨機応変に＝四字熟語）"]),
("9. 〜するほど・〜するあまり：đến mức / đến nỗi",
 "程度が大きいことを強調する。",
 ["第15問 …nhiều đến mức học sinh muốn khóc.",
  "第16問 …mải chơi game đến nỗi quên ăn.　第22問 …cảm động đến mức rơi nước mắt."]),
("10. 〜するために：để + 動詞",
 "目的を表す。",
 ["第17問 …xếp hàng … để làm thủ tục."]),
("11. 〜だけ…する（同じ量）：bao nhiêu … bấy nhiêu",
 "前の量と同じだけ後ろもする。",
 ["第18問 Chị muốn mua bao nhiêu thì tôi sẽ bán cho chị bấy nhiêu."]),
("12. 〜のとおりに…する：sao … vậy",
 "言われたとおりに行う。",
 ["第19問 Giám đốc yêu cầu sao thì tôi làm vậy."]),
("13. 〜すると決まって…：cứ … là …",
 "あることが起きると必ず次が起きる。",
 ["第20問 Cứ mưa to là đường phố Hà Nội bị ngập."]),
("14. 〜させる（引き起こす）：khiến / làm / làm cho",
 "原因が結果を引き起こす言い方。",
 ["第21問 Tiếng còi xe tải khiến tôi giật mình.",
  "第22問 …làm tôi cảm động…　第23問 …đã làm cho ngành du lịch tê liệt…"]),
("15. もし〜なら…（すべき）：nếu … thì (nên) …",
 "条件と、それに対する助言・結果。",
 ["第24問 Nếu anh muốn giải quyết nhanh chóng thì nên thuê luật sư."]),
("16. 自分で〜する：tự … lấy",
 "自分の力で行うことを強める。",
 ["第25問 Tôi tự may quần áo lấy."]),
("17. 〜と言える：có thể nói",
 "「〜と言ってよい」という前置き。",
 ["第32問 Có thể nói núi Phú Sĩ là biểu tượng của Nhật Bản.",
  "第33問 Có thể nói Covid-19 đã thay đổi thói quen người tiêu dùng."]),
("18. その他",
 "",
 ["trước khi ＋動詞 ＝ 〜する前に（第10問）",
  "chưa bao giờ ＝ 一度も〜ない（第28問）／ không bao giờ ＝ 決して〜ない（第29問）"]),
]
for title,desc,exs in grammar:
    para(doc,title,size=11.5,bold=True,after=2,before=4)
    if desc: para(doc,desc,size=10.5,after=2)
    for e in exs: para(doc,"・"+e,size=10.5,after=1)

# ===== 単語 =====
doc.add_page_break()
head(doc,"Ⅱ. 単語（漢越語つき）")
para(doc,"「漢越語」列＝漢字に対応するベトナム語。— は固有語（漢越語でない）。"
         "（diễn演）のように一部だけが漢越語の場合はその部分を示す。",
     size=9.5,color=(0x66,0x66,0x66),after=8)

vocab=[
("phản đối","反対する","反対"),
("cưới","結婚する（娶る）","—"),
("cố gắng hết sức","全力を尽くす","—"),
("thay đổi","変える・変わる","—"),
("tình hình","情勢・状況","情形"),
("phát minh","発明する","発明"),
("bản báo cáo","報告書","報告（báo cáo）／本（bản）"),
("diễn viên","俳優","演員"),
("nhà văn","作家","文（văn）"),
("hồ (nước)","湖","湖（hồ）"),
("căn bếp","台所","—"),
("nước mắm","魚醤（ヌクマム）","—"),
("suy nghĩ","考える","推（suy）"),
("trọng lượng","重さ","重量"),
("kích thước","大きさ・寸法","—"),
("giá tiền","値段","価（giá）"),
("câu trả lời","答え","—"),
("cách suy nghĩ","考え方","格（cách）・推（suy）"),
("bài tập về nhà","宿題","習（tập）"),
("khóc","泣く","—"),
("quên","忘れる","—"),
("hành khách","乗客","行客"),
("xếp hàng","列に並ぶ","—"),
("sân bay","空港","—"),
("làm thủ tục","手続きをする","手続（thủ tục）"),
("giám đốc","社長・部長","監督"),
("yêu cầu","要求する","要求"),
("mưa to","大雨","—"),
("đường phố","通り","—"),
("bị ngập","水浸しになる","—"),
("giật mình","びっくりする","—"),
("cảm động","感動する","感動"),
("rơi nước mắt","涙を流す","—"),
("ngành du lịch","観光業","遊歴（du lịch）"),
("tê liệt","麻痺する","—"),
("hoàn toàn","完全に","完全"),
("giải quyết","解決する","解決"),
("thuê","雇う・借りる","—"),
("luật sư","弁護士","律師"),
("đồng ý","同意する","同意"),
("đi du học","留学する","遊学（du học）"),
("thời tiết","天気","時節"),
("trận đấu","試合","—"),
("diễn ra","行われる","演（diễn）"),
("nghèo","貧しい","—"),
("ăn trộm","盗む","—"),
("tặng quà","贈り物をする","贈（tặng）"),
("bận rộn","忙しい","—"),
("rảnh rỗi","暇","—"),
("gọi điện","電話する","電（điện）"),
("hằng ngày","毎日","恒（hằng）"),
("biểu tượng","象徴","表象"),
("thói quen","習慣","—"),
("người tiêu dùng","消費者","消（tiêu）"),
("thường xuyên","頻繁に","常（thường）"),
("đau dạ dày","胃が痛い","—"),
("lười học","勉強を怠ける","学（học）"),
("thi lại","追試・再試験","試（thi）"),
]
tbl=doc.add_table(rows=1,cols=3); tbl.style='Table Grid'
for c,t in zip(tbl.rows[0].cells,["ベトナム語","意味","漢越語（漢字）"]):
    r=c.paragraphs[0].add_run(t); setf(r,size=10.5,bold=True)
for vi,jp,hv in vocab:
    cells=tbl.add_row().cells
    setf(cells[0].paragraphs[0].add_run(vi),size=11)
    setf(cells[1].paragraphs[0].add_run(jp),size=11)
    setf(cells[2].paragraphs[0].add_run(hv),size=11)

out=r"C:\Users\baru1\Desktop\ベトナム語\Huong並び替え_文法と単語_漢越語つき.docx"
doc.save(out)
print("SAVED:",out)
print("文法",len(grammar),"単語",len(vocab))
