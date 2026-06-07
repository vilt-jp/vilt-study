# -*- coding: utf-8 -*-
"""Thanh Nien記事「AIで宿題」の4級やさしい版Word作成。
構成: ①やさしい越語原文 ②読解5問 ③(改ページ)日本語訳 ④4級単語・文法。
フォントMeiryo UI(latin/ea/cs)。元の上級版とは別ファイル。"""
from docx import Document
from docx.shared import Pt, RGBColor
from docx.oxml.ns import qn
from docx.enum.text import WD_ALIGN_PARAGRAPH

FONT="Meiryo UI"
def setf(run,size=None,bold=None,color=None):
    run.font.name=FONT
    rPr=run._element.get_or_add_rPr()
    rF=rPr.find(qn('w:rFonts'))
    if rF is None:
        rF=rPr.makeelement(qn('w:rFonts'),{}); rPr.insert(0,rF)
    for a in ('w:ascii','w:hAnsi','w:eastAsia','w:cs'): rF.set(qn(a),FONT)
    if size is not None: run.font.size=Pt(size)
    if bold is not None: run.font.bold=bold
    if color is not None: run.font.color.rgb=RGBColor(*color)

def para(doc,text="",size=11,bold=False,color=None,align=None,after=6):
    p=doc.add_paragraph()
    if align is not None: p.alignment=align
    p.paragraph_format.space_after=Pt(after)
    r=p.add_run(text); setf(r,size=size,bold=bold,color=color); return p

def head(doc,text,size=14):
    return para(doc,text,size=size,bold=True,color=(0x1F,0x49,0x7D),after=8)

doc=Document()
st=doc.styles['Normal']; st.font.name=FONT; st.element.rPr.rFonts.set(qn('w:eastAsia'),FONT)

para(doc,"Hơn một nửa học sinh dùng AI để làm bài tập",size=16,bold=True,
     color=(0x1F,0x49,0x7D),align=WD_ALIGN_PARAGRAPH.CENTER,after=4)
para(doc,"【4級やさしい版】 元記事: Báo Thanh Niên（やさしい言葉に書き直したもの）",
     size=9.5,color=(0x66,0x66,0x66),align=WD_ALIGN_PARAGRAPH.CENTER,after=10)

# ① やさしい越語
head(doc,"Ⅰ. Bài đọc (やさしいベトナム語)")
vn=[
"Một trung tâm nghiên cứu ở Mỹ tên là Pew đã làm một cuộc khảo sát về việc học sinh dùng AI.",
"Kết quả cho thấy 54% học sinh từ 13 đến 17 tuổi đã dùng AI như ChatGPT để làm bài tập. 57% dùng AI để tìm thông tin. 47% dùng để giải trí. 42% dùng để tóm tắt bài.",
"Nhiều bố mẹ cũng đồng ý cho con dùng AI. 58% bố mẹ thấy không sao khi con dùng AI để làm bài. Nhưng nếu con dùng AI để nói chuyện hay xin lời khuyên thì ít bố mẹ đồng ý hơn, chỉ 28% và 18%.",
"Cuộc khảo sát hỏi 1.458 cặp học sinh và bố mẹ, từ ngày 25 tháng 9 đến ngày 9 tháng 10 năm 2025. 44% học sinh nói chỉ dùng AI cho một ít bài tập. 10% nói dùng AI cho gần hết bài tập.",
"Một điều đáng chú ý là 59% học sinh nghĩ rằng ở trường mình, nhiều bạn dùng AI để gian lận. Trong nhóm học sinh có dùng AI, 76% nghĩ rằng bạn cùng trường dùng AI để gian lận.",
"Một học sinh nói: “Mọi người dùng AI quá nhiều để làm bài và hỏi những câu dễ.”",
"Có người nói trường nên dạy học sinh cách dùng AI cho đúng. Nhưng cũng có người lo rằng AI làm cho học sinh lười suy nghĩ và dễ gian lận hơn.",
"Một nghiên cứu của trường đại học Cambridge và công ty Microsoft cho thấy: học sinh tự ghi chú mà không dùng AI thì hiểu bài tốt hơn học sinh dùng AI.",
"Một nhà tâm lý học khuyên bố mẹ nên chú ý khi con dùng AI quá nhiều, không chịu suy nghĩ, hay có vẻ buồn.",
]
for t in vn: para(doc,t,size=11,after=6)

# ② 読解問題
doc.add_paragraph()
head(doc,"Ⅱ. Câu hỏi đọc hiểu (読解問題・5問)")
para(doc,"Đọc bài trên và chọn đáp án đúng. (本文を読み、正しい答えを選びなさい)",
     size=10,color=(0x66,0x66,0x66),after=8)
qs=[
("Câu 1. Bao nhiêu phần trăm học sinh đã dùng AI để làm bài tập?",
 ["A. 42%","B. 47%","C. 54%","D. 57%"]),
("Câu 2. Bao nhiêu phần trăm bố mẹ thấy không sao khi con dùng AI để làm bài?",
 ["A. 28%","B. 18%","C. 44%","D. 58%"]),
("Câu 3. Bao nhiêu phần trăm học sinh nghĩ rằng ở trường mình nhiều bạn dùng AI để gian lận?",
 ["A. 10%","B. 59%","C. 76%","D. 54%"]),
("Câu 4. Theo nghiên cứu của Cambridge, ai hiểu bài tốt hơn?",
 ["A. Học sinh dùng AI","B. Học sinh tự ghi chú mà không dùng AI",
  "C. Bố mẹ","D. Giáo viên"]),
("Câu 5. Nhà tâm lý học khuyên bố mẹ làm gì?",
 ["A. Cho con dùng AI nhiều hơn","B. Cấm con đi học",
  "C. Chú ý khi con dùng AI quá nhiều","D. Mua máy tính mới"]),
]
for q,opts in qs:
    para(doc,q,size=11,bold=True,after=2)
    for o in opts: para(doc,"   "+o,size=11,after=1)
    doc.add_paragraph()
para(doc,"【解答】 Câu 1: C  /  Câu 2: D  /  Câu 3: B  /  Câu 4: B  /  Câu 5: C",
     size=11,bold=True,color=(0xC0,0,0),after=6)

# ③ 日本語訳
doc.add_page_break()
head(doc,"Ⅲ. 日本語訳")
para(doc,"見出し: 「半分以上の生徒が、宿題にAIを使う」",size=11,bold=True,after=8)
jp=[
"Pewという名前の、アメリカにある研究センターが、生徒のAI利用について調査をした。",
"その結果、13歳から17歳の生徒の54%が、ChatGPTのようなAIを宿題をするために使ったことが分かった。57%は情報を探すために使う。47%は娯楽のために、42%は文章を要約するために使う。",
"多くの親も、子どもがAIを使うことに同意している。58%の親は、子どもが宿題にAIを使っても問題ないと考えている。ただし、おしゃべりや助言を求めるためにAIを使う場合は、同意する親は少なく、それぞれ28%と18%だけだった。",
"この調査は、生徒と親1,458組に、2025年9月25日から10月9日まで聞いた。44%の生徒は、少しの宿題だけにAIを使うと答えた。10%は、ほとんどの宿題にAIを使うと答えた。",
"注目すべきことに、59%の生徒が、自分の学校では多くの友達がAIを使って不正をしていると思っている。AIを使う生徒の中では、76%が、同じ学校の友達がAIで不正をしていると思っている。",
"ある生徒はこう言った。「みんな、宿題をするためや簡単なことを聞くために、AIを使いすぎている。」",
"学校は生徒にAIの正しい使い方を教えるべきだ、と言う人もいる。しかし、AIは生徒を考えない人にし、不正もしやすくする、と心配する人もいる。",
"ケンブリッジ大学とMicrosoft社の研究では、AIを使わずに自分でメモを取った生徒のほうが、AIを使った生徒よりも内容をよく理解できた、と分かった。",
"ある心理学者は、子どもがAIを使いすぎる、自分で考えようとしない、悲しそうにしている時は、親が気をつけるとよいと勧めている。",
]
for t in jp: para(doc,t,size=11,after=6)

# ④ 4級単語・文法
doc.add_paragraph()
head(doc,"Ⅳ. 4級の単語・文法")
para(doc,"▼ 単語",size=12,bold=True,color=(0x1F,0x49,0x7D),after=4)
vocab=[
("học sinh","生徒"),
("bài tập","宿題・課題"),
("tìm thông tin","情報を探す"),
("giải trí","娯楽・気晴らし"),
("tóm tắt","要約する"),
("bố mẹ","両親"),
("đồng ý","同意する"),
("lời khuyên","助言"),
("gian lận","不正・カンニング"),
("đáng chú ý","注目すべき"),
("suy nghĩ","考える"),
("ghi chú","メモを取る"),
("hiểu bài","内容を理解する"),
("khuyên","忠告する・勧める"),
("chú ý","注意する・気をつける"),
("buồn","悲しい"),
("kết quả","結果"),
("cuộc khảo sát","調査"),
]
tbl=doc.add_table(rows=1,cols=2); tbl.style='Table Grid'
for c,t in zip(tbl.rows[0].cells,["ベトナム語","意味"]):
    setf(c.paragraphs[0].add_run(t),size=10.5,bold=True)
for vi,jp_ in vocab:
    cells=tbl.add_row().cells
    setf(cells[0].paragraphs[0].add_run(vi),size=11)
    setf(cells[1].paragraphs[0].add_run(jp_),size=11)

doc.add_paragraph()
para(doc,"▼ 文法（やさしい言い方）",size=12,bold=True,color=(0x1F,0x49,0x7D),after=4)
gram=[
("cho thấy","「〜を示す・〜だと分かる」。例: Kết quả cho thấy…（結果は〜を示す）"),
("như","「〜のような」。例: AI như ChatGPT（ChatGPTのようなAI）"),
("để + 動詞","「〜するために」。例: dùng AI để làm bài（宿題をするためにAIを使う）"),
("nếu … thì …","「もし〜なら…」。例: nếu con dùng AI thì…（もし子がAIを使うなら）"),
("nhưng","「しかし」。前と逆のことを続ける。"),
("cũng","「〜も」。例: bố mẹ cũng đồng ý（親も同意する）"),
("nên","「〜したほうがよい・〜すべき」。例: trường nên dạy（学校は教えるべき）"),
("nghĩ rằng / lo rằng","「〜と思う／〜と心配する」。rằng のあとに内容の文。"),
("… mà không …","「〜なのに…ない／〜せずに」。例: tự ghi chú mà không dùng AI（AIを使わずに自分でメモする）"),
("hơn","比較「〜より」。例: hiểu bài tốt hơn（よりよく理解する）"),
]
for pat,exp in gram:
    para(doc,"・"+pat,size=11,bold=True,after=1)
    para(doc,"    "+exp,size=10.5,after=6)

out=r"C:\Users\baru1\Desktop\ベトナム語\ベトナム語検定\4級\Thanh Nien_AIで宿題_4級やさしい版.docx"
doc.save(out)
print("SAVED:",out)
print("越語段落",len(vn),"設問",len(qs),"訳",len(jp),"単語",len(vocab),"文法",len(gram))
