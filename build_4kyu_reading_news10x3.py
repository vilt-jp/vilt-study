# -*- coding: utf-8 -*-
"""4級読解 時事対策（拡充版v2）。長文10本×各3問=30問。本文をv1の約1.5倍に延長。
構成: ①やさしい越語本文10本 ②読解30問+解答 ③(改ページ)日本語訳10本 ④4級単語・文法。
フォントMeiryo UI(latin/ea/cs)。練習用(傾向ベースの作文・実在記事ではない)。"""
from docx import Document
from docx.shared import Pt, RGBColor
from docx.oxml.ns import qn
from docx.enum.text import WD_ALIGN_PARAGRAPH

FONT="Meiryo UI"
def setf(run,size=None,bold=None,color=None):
    run.font.name=FONT
    rPr=run._element.get_or_add_rPr()
    rF=rPr.find(qn('w:rFonts'))
    if rF is None:
        rF=rPr.makeelement(qn('w:rFonts'),{}); rPr.insert(0,rF)
    for a in ('w:ascii','w:hAnsi','w:eastAsia','w:cs'): rF.set(qn(a),FONT)
    if size is not None: run.font.size=Pt(size)
    if bold is not None: run.font.bold=bold
    if color is not None: run.font.color.rgb=RGBColor(*color)
def para(doc,text="",size=11,bold=False,color=None,align=None,after=6):
    p=doc.add_paragraph()
    if align is not None: p.alignment=align
    p.paragraph_format.space_after=Pt(after)
    setf(p.add_run(text),size=size,bold=bold,color=color); return p
def head(doc,text,size=14):
    return para(doc,text,size=size,bold=True,color=(0x1F,0x49,0x7D),after=8)

P=[
("【1 経済・社会】 Thanh toán không dùng tiền mặt（キャッシュレス決済）",
"Hiện nay, ngày càng nhiều người Việt Nam thanh toán không dùng tiền mặt. Khi đi mua đồ ở siêu thị, quán cà phê, nhà hàng hay thậm chí ở chợ, nhiều người chỉ cần lấy điện thoại ra và quét mã QR. Sau vài giây, việc trả tiền đã xong. Cách này vừa nhanh vừa tiện, lại không cần mang theo nhiều tiền lẻ. Người dùng cũng không lo bị nhận nhầm tiền giả hay phải chờ người bán trả lại tiền thừa. Các ngân hàng và công ty công nghệ đã làm ra nhiều ứng dụng để việc trả tiền dễ hơn. Nhà nước cũng khuyến khích người dân dùng cách thanh toán mới này. Nhờ vậy, nhiều cửa hàng nhỏ, quán ăn ven đường và cả người bán hàng ở chợ bây giờ cũng dán mã QR ở quầy. Một số người còn dùng điện thoại để trả tiền điện, tiền nước và mua vé xe. Tuy nhiên, không phải ai cũng thích cách này. Một số người lớn tuổi vẫn quen dùng tiền mặt và thấy khó khi dùng điện thoại để trả tiền. Ngoài ra, một số người lo rằng nếu điện thoại hết pin hoặc mất mạng thì sẽ không trả được tiền. Cũng có người sợ bị kẻ xấu lấy mất tiền trong tài khoản nếu không cẩn thận. Vì vậy, các chuyên gia khuyên mọi người nên đặt mật khẩu khó và không bấm vào những đường link lạ. Dù còn vài khó khăn, nhiều chuyên gia tin rằng trong tương lai, cách thanh toán này sẽ ngày càng phổ biến hơn ở Việt Nam.",
"今、ますます多くのベトナム人が現金を使わずに支払いをしている。スーパー、喫茶店、レストラン、さらには市場で買い物をするとき、多くの人は電話を取り出してQRコードを読み取るだけでよい。数秒で支払いが終わる。この方法は速くて便利で、しかも小銭をたくさん持ち歩く必要もない。利用者は、偽札を受け取ったり、店がおつりを返すのを待ったりする心配もない。銀行や技術系の会社は、支払いをより簡単にするためのアプリをたくさん作ってきた。国も、この新しい支払い方法を使うよう国民に勧めている。そのおかげで、小さな店、道ばたの食堂、市場の売り手までも、今ではレジにQRコードを貼っている。電話で電気代や水道代を払ったり、乗車券を買ったりする人もいる。ただし、誰もがこの方法を好むわけではない。一部の年配の人は今も現金に慣れていて、電話で支払うのを難しく感じる。さらに、電話の電池が切れたり通信ができなくなったりしたら支払えない、と心配する人もいる。注意しないと悪い人に口座のお金を取られるのを怖がる人もいる。だから専門家は、難しい暗証番号を設定し、知らないリンクを押さないよう勧めている。まだいくつかの難しさはあるが、多くの専門家は、将来この支払い方法はベトナムでますます広まると信じている。",
[("Câu 1. Người ta trả tiền bằng cách quét mã QR ở đâu?",
  ["A. Chỉ ở ngân hàng","B. Ở siêu thị, quán cà phê, nhà hàng, chợ","C. Chỉ ở nhà","D. Chỉ ở sân bay"],1),
 ("Câu 2. Vì sao một số người lớn tuổi thấy khó?",
  ["A. Vì họ không có tiền","B. Vì cửa hàng không nhận","C. Vì vẫn quen dùng tiền mặt","D. Vì giá quá cao"],2),
 ("Câu 3. Người ta lo lắng điều gì về cách trả tiền này?",
  ["A. Điện thoại hết pin hoặc mất mạng thì không trả được","B. Tiền lẻ quá nhiều","C. Cửa hàng đóng cửa","D. Không có ứng dụng"],0)]),

("【2 経済・社会】 Mua sắm trực tuyến（オンラインショッピング）",
"Vài năm gần đây, mua sắm trực tuyến trở nên rất phổ biến ở Việt Nam. Người ta có thể mua quần áo, đồ ăn, sách và nhiều thứ khác chỉ bằng điện thoại. Khách hàng không phải đi ra cửa hàng mà vẫn chọn được hàng ngay tại nhà. Trên mạng có rất nhiều loại hàng để chọn, từ hàng trong nước đến hàng nước ngoài. Sau khi đặt hàng, người giao hàng sẽ mang đồ đến tận nơi, thường chỉ trong vài ngày. Giá trên mạng đôi khi rẻ hơn ở cửa hàng, và có nhiều chương trình giảm giá vào những ngày đặc biệt. Nhiều người bán còn phát trực tiếp trên mạng để giới thiệu sản phẩm cho khách xem. Vì những lý do đó, nhiều người trẻ rất thích mua sắm theo cách này. Tuy nhiên, mua trên mạng cũng có điểm không tốt. Khách không được nhìn và thử hàng trực tiếp, nên đôi khi hàng nhận được không giống như trong ảnh. Có người còn mua phải hàng kém chất lượng hoặc hàng giả. Khi muốn đổi hay trả lại hàng, đôi khi cũng mất nhiều thời gian. Vì thế, các chuyên gia khuyên người mua nên chọn cửa hàng có uy tín và đọc kỹ ý kiến của khách hàng khác trước khi đặt. Người mua cũng nên kiểm tra kỹ hàng ngay khi nhận. Nếu mua một cách thông minh, mua sắm trực tuyến sẽ vừa tiện vừa tiết kiệm.",
"ここ数年、ベトナムではオンラインショッピングがとても一般的になった。人々は服、食べ物、本、その他多くの物を電話だけで買うことができる。客は店に出かけなくても、家でそのまま商品を選べる。ネット上には国内の品から外国の品まで、とても多くの種類がある。注文したあと、配達の人がその場所まで品物を持ってきてくれて、たいてい数日でとどく。ネットの値段は店より安いこともあり、特別な日には割引もたくさんある。売り手の多くは、客に見せるためにネットで生放送をして商品を紹介する。こうした理由から、多くの若者はこの買い方をとても好む。ただし、ネットで買うことには良くない点もある。客は直接見たり試したりできないので、受け取った品物が写真と違うこともある。質の悪い品物や偽物を買ってしまう人もいる。交換や返品をしたいとき、時間がかかることもある。だから専門家は、注文する前に信用のある店を選び、ほかの客の意見をよく読むよう勧めている。受け取ったらすぐに品物をよく確かめるのもよい。賢く買えば、オンラインショッピングは便利で節約にもなる。",
[("Câu 4. Mua sắm trực tuyến có lợi gì?",
  ["A. Phải đi ra cửa hàng","B. Chọn hàng ở nhà, có người giao đến tận nơi","C. Hàng luôn đắt hơn","D. Không có giảm giá"],1),
 ("Câu 5. Điểm không tốt của mua trên mạng là gì?",
  ["A. Không được nhìn và thử hàng trực tiếp","B. Giao hàng quá nhanh","C. Không có ứng dụng","D. Quá nhiều cửa hàng"],0),
 ("Câu 6. Chuyên gia khuyên người mua nên làm gì?",
  ["A. Mua thật nhiều","B. Không đọc ý kiến người khác","C. Chọn cửa hàng có uy tín và đọc kỹ ý kiến khách","D. Chỉ mua hàng đắt"],2)]),

("【3 くらし・働き方】 Làm việc ở nhà（在宅勤務）",
"Sau dịch Covid-19, cách làm việc của nhiều người đã thay đổi. Trước đây, hầu hết nhân viên phải đến công ty mỗi ngày, có người mất một, hai tiếng chỉ để đi lại. Bây giờ, nhiều công ty cho phép nhân viên làm việc ở nhà. Nhờ máy tính và mạng internet, nhân viên vẫn có thể họp trực tuyến và gửi tài liệu cho nhau dù ở xa. Có nhiều ứng dụng giúp mọi người nói chuyện và làm việc cùng nhau dễ dàng. Làm việc ở nhà giúp tiết kiệm thời gian và tiền đi lại. Nhân viên cũng có thêm thời gian cho gia đình và được nghỉ ngơi nhiều hơn. Một số người nói rằng khi làm ở nhà, họ thấy thoải mái và làm việc tốt hơn. Tuy nhiên, cách làm việc này cũng có khó khăn. Ở nhà, đôi khi khó tập trung vì có nhiều việc khác như con cái hay việc nhà. Một số người cảm thấy cô đơn vì ít gặp đồng nghiệp và ít được nói chuyện trực tiếp. Cũng có người làm việc quá nhiều vì không phân biệt rõ giờ làm và giờ nghỉ. Vì vậy, hiện nay nhiều công ty chọn cách kết hợp: nhân viên vừa làm ở nhà vài ngày, vừa lên công ty vài ngày trong tuần. Cách này giúp nhân viên vừa thoải mái vừa vẫn gặp được đồng nghiệp. Nhiều người cho rằng đây là cách tốt cho cả công ty và nhân viên.",
"新型コロナの流行のあと、多くの人の働き方が変わった。以前は、ほとんどの従業員が毎日会社に来なければならず、通勤だけで一、二時間かかる人もいた。今は、多くの会社が従業員に在宅勤務を許している。パソコンとインターネットのおかげで、遠くにいてもオンライン会議をしたり、書類を送り合ったりできる。みんなが話したり一緒に働いたりしやすくするアプリもたくさんある。在宅勤務は時間と通勤のお金を節約できる。従業員は家族のための時間も増え、より休めるようになる。家で働くと、気持ちが楽で仕事もよくできると言う人もいる。ただし、この働き方にも難しさがある。家では、子どもや家事などほかの用事が多くて集中しにくいことがある。同僚にあまり会わず、直接話す機会も少ないので、さびしいと感じる人もいる。仕事の時間と休みの時間をはっきり分けられず、働きすぎる人もいる。そのため、今は多くの会社が組み合わせの方法を選んでいる。週のうち数日は家で働き、数日は会社に行くというものだ。この方法だと、従業員は楽でありながら同僚にも会える。これは会社にも従業員にも良い方法だと、多くの人が考えている。",
[("Câu 7. Nhờ đâu mà nhân viên có thể làm việc ở nhà?",
  ["A. Nhờ máy tính và mạng internet","B. Nhờ xe máy","C. Nhờ tiền mặt","D. Nhờ thời tiết"],0),
 ("Câu 8. Khó khăn khi làm việc ở nhà là gì?",
  ["A. Không có nhà","B. Khó tập trung và cảm thấy cô đơn","C. Không có gia đình","D. Lương quá cao"],1),
 ("Câu 9. Cách kết hợp mà nhiều công ty chọn là gì?",
  ["A. Chỉ làm ở nhà","B. Chỉ lên công ty","C. Vừa làm ở nhà vừa lên công ty trong tuần","D. Nghỉ cả tuần"],2)]),

("【4 くらし・働き方】 Tập thể dục（運動・健康ブーム）",
"Gần đây, ngày càng nhiều người Việt Nam quan tâm đến sức khỏe. Vào buổi sáng sớm hoặc buổi chiều, ở các công viên có rất nhiều người đi bộ, chạy bộ và tập thể dục. Nhiều người trẻ còn đến phòng tập để tập với máy, có người tập yoga hoặc bơi lội. Một số người tham gia các câu lạc bộ thể thao để tập cùng bạn bè cho vui. Họ cho rằng tập thể dục đều đặn giúp cơ thể khỏe mạnh, ít bị bệnh và tinh thần thoải mái hơn. Nhiều người dùng điện thoại hoặc đồng hồ thông minh để đếm số bước chân mỗi ngày. Ngoài ra, nhiều người cũng chú ý hơn đến việc ăn uống, ăn nhiều rau, ít đồ ngọt và uống đủ nước. Một số người còn nấu ăn ở nhà thay vì ăn ngoài để bảo vệ sức khỏe. Tuy nhiên, các bác sĩ nhắc rằng không nên tập quá sức. Nếu tập sai cách hoặc tập quá nhiều, người ta có thể bị đau hoặc bị thương. Vì vậy, trước khi tập nặng, mọi người nên hỏi ý kiến bác sĩ và bắt đầu từ từ. Người mới tập nên chọn bài tập nhẹ rồi tăng dần. Các bác sĩ cũng nói rằng sống khỏe không chỉ là tập thể dục mà còn là ngủ đủ giấc và giữ tinh thần vui vẻ. Khi cơ thể và tinh thần đều khỏe, người ta sẽ làm việc và học tập tốt hơn.",
"最近、ますます多くのベトナム人が健康に関心を持っている。早朝や夕方、公園にはたくさんの人が散歩したり、走ったり、運動したりしている。多くの若者は、機械で運動するためにジムにも行き、ヨガをしたり泳いだりする人もいる。楽しむために友達と一緒に運動しようと、スポーツのクラブに入る人もいる。彼らは、規則正しく運動すると体が丈夫になり、病気になりにくく、気持ちもより楽になると考えている。多くの人は毎日の歩数を数えるために電話やスマートウォッチを使う。さらに、多くの人が食事にも前より気をつけ、野菜を多く、甘い物を少なくし、水を十分に飲んでいる。健康を守るために外食の代わりに家で料理する人もいる。ただし、医者は、運動しすぎないようにと注意している。間違ったやり方や、やりすぎで、痛めたりけがをしたりすることがある。だから、きつい運動の前には医者の意見を聞き、少しずつ始めるとよい。始めたばかりの人は軽い運動を選び、だんだん増やすのがよい。医者はまた、健康に暮らすことは運動だけでなく、十分に眠り、気持ちを明るく保つことでもあると言う。体と心の両方が健康なら、仕事も勉強もよりうまくいく。",
[("Câu 10. Buổi sáng và buổi chiều, ở công viên người ta làm gì?",
  ["A. Mua sắm","B. Đi bộ, chạy bộ và tập thể dục","C. Làm việc","D. Ngủ"],1),
 ("Câu 11. Tập thể dục đều đặn có lợi gì?",
  ["A. Tốn nhiều tiền","B. Cơ thể khỏe mạnh và tinh thần thoải mái","C. Khó ngủ hơn","D. Hay bị bệnh hơn"],1),
 ("Câu 12. Bác sĩ nhắc mọi người điều gì?",
  ["A. Nên tập càng nặng càng tốt","B. Không cần hỏi ai","C. Không nên tập quá sức, nên bắt đầu từ từ","D. Không nên ăn rau"],2)]),

("【5 文化・行事・観光】 Du lịch（観光客の回復）",
"Năm nay, rất nhiều khách du lịch nước ngoài đến Việt Nam. Các thành phố như Đà Nẵng, Hà Nội, Hội An và Nha Trang luôn đông khách. Nhiều khách đến từ Hàn Quốc, Nhật Bản và các nước châu Âu. Họ thích ăn các món ăn Việt Nam như phở và bún chả, đi thăm phố cổ và mua đồ lưu niệm. Nhiều người còn thích đi biển, ngắm cảnh đẹp và xem các lễ hội truyền thống. Việt Nam có nhiều nơi đẹp, thời tiết ấm và giá cả không quá đắt, nên được nhiều khách yêu thích. Nhờ có nhiều khách, các khách sạn, nhà hàng và cửa hàng đều bán được nhiều hơn. Nhiều người dân cũng có thêm việc làm và thêm thu nhập. Ngành du lịch vì thế phát triển trở lại sau mấy năm khó khăn. Để đón khách tốt hơn, nhiều nơi đã sửa lại đường, khách sạn và làm thêm các tour mới. Một số người dân địa phương cũng học thêm tiếng nước ngoài để nói chuyện với khách. Chính phủ cũng giúp khách dễ xin giấy tờ vào Việt Nam hơn trước. Tuy nhiên, khi quá đông khách, một số nơi đẹp có thể bị bẩn hoặc hư hỏng, và giá cả đôi khi tăng cao. Vì vậy, mọi người cần giữ gìn cảnh đẹp và môi trường. Nếu làm tốt, du lịch sẽ phát triển lâu dài và mang lại lợi ích cho nhiều người.",
"今年は、とても多くの外国人観光客がベトナムに来ている。ダナン、ハノイ、ホイアン、ニャチャンといった都市はいつも観光客で混んでいる。多くの客は韓国、日本、ヨーロッパの国々から来る。彼らはフォーやブンチャーなどのベトナム料理を食べたり、旧市街を訪れたり、おみやげを買ったりするのが好きだ。海に行ったり、美しい景色を眺めたり、伝統的な祭りを見たりするのが好きな人も多い。ベトナムには美しい場所が多く、気候は暖かく、値段も高すぎないので、多くの客に好まれている。客が多いおかげで、ホテル、レストラン、店はどこも以前より多く売れている。多くの住民も仕事と収入が増える。観光業は、数年の苦しい時期のあと、そのため再び発展している。よりよく客を迎えるため、多くの場所が道路やホテルを直し、新しいツアーも作った。一部の地元の人は、客と話すために外国語も学んでいる。政府も、客が以前より簡単にベトナム入国の書類を取れるようにしている。ただし、客が多すぎると、美しい場所が汚れたり傷んだりし、値段が高くなることもある。だから、みんなが景色や環境を大切にする必要がある。うまくやれば、観光は長く発展し、多くの人に利益をもたらす。",
[("Câu 13. Khách du lịch nước ngoài thích làm gì ở Việt Nam?",
  ["A. Ăn món Việt, thăm phố cổ, mua đồ lưu niệm","B. Chỉ ở trong khách sạn","C. Làm việc","D. Học đại học"],0),
 ("Câu 14. Nhờ có nhiều khách, điều gì xảy ra?",
  ["A. Cửa hàng đóng cửa","B. Khách sạn, nhà hàng, cửa hàng bán được nhiều hơn","C. Đường bị hỏng hết","D. Không ai có việc làm"],1),
 ("Câu 15. Khi quá đông khách, có thể có vấn đề gì?",
  ["A. Không có vấn đề gì","B. Khách không đến nữa","C. Nơi đẹp có thể bị bẩn hoặc hư hỏng","D. Giá xuống thấp"],2)]),

("【6 文化・行事・観光】 Tết（旧正月の過ごし方の変化）",
"Tết là ngày lễ quan trọng nhất của người Việt Nam. Vào dịp Tết, mọi người được nghỉ dài và thường về quê thăm gia đình. Họ dọn nhà, mua hoa, gói bánh chưng, nấu các món ăn truyền thống và mừng tuổi cho trẻ em. Tết cũng là lúc mọi người đi thăm họ hàng, bạn bè và chúc nhau những điều tốt đẹp cho năm mới. Trẻ em rất thích Tết vì được nghỉ học, được mặc quần áo mới và được nhận tiền mừng tuổi. Tuy nhiên, gần đây cách đón Tết của một số gia đình đã thay đổi. Thay vì chỉ ở nhà, nhiều gia đình chọn đi du lịch trong nước hoặc ra nước ngoài trong dịp Tết. Có người nói rằng đi chơi giúp cả nhà nghỉ ngơi, thay đổi không khí và có thêm kỷ niệm vui. Cũng có người mua sẵn đồ ăn làm sẵn để không phải nấu nhiều như trước. Nhưng cũng có người cho rằng Tết nên ở bên ông bà, bố mẹ thì mới có ý nghĩa. Họ lo rằng các phong tục đẹp ngày Tết sẽ dần bị quên. Dù mỗi nhà mỗi khác, hầu hết người Việt vẫn xem Tết là dịp để sum họp và nghỉ ngơi sau một năm làm việc vất vả. Nhiều bạn trẻ tuy đi làm xa nhưng Tết nào cũng cố gắng về nhà để được ở bên gia đình.",
"テトはベトナム人にとって最も大切な祝日だ。テトの時期、人々は長く休み、たいてい故郷に帰って家族を訪ねる。家を掃除し、花を買い、バインチュンを包み、伝統的な料理を作り、子どもにお年玉をあげる。テトはまた、親戚や友人を訪ね、新年に良いことを願い合うときでもある。子どもは、学校が休みで、新しい服を着られ、お年玉をもらえるので、テトが大好きだ。ただし、最近は一部の家庭のテトの過ごし方が変わってきた。家にいるだけでなく、テトの時期に国内や外国へ旅行することを選ぶ家庭も多い。出かけることで家族みんなが休め、気分を変えられ、楽しい思い出も増えると言う人もいる。前ほどたくさん料理しなくてよいように、出来合いの料理を買っておく人もいる。しかし、テトは祖父母や両親のそばにいてこそ意味がある、と考える人もいる。彼らは、テトの美しい風習がだんだん忘れられることを心配している。家ごとに違うとはいえ、ほとんどのベトナム人は今もテトを、一年の大変な仕事のあとに集まり休む機会だと考えている。多くの若者は遠くで働いていても、家族のそばにいるために、どのテトもがんばって家に帰る。",
[("Câu 16. Vào dịp Tết, người Việt thường làm gì?",
  ["A. Đi làm bình thường","B. Về quê thăm gia đình, dọn nhà, nấu món truyền thống","C. Không nghỉ","D. Bán nhà"],1),
 ("Câu 17. Gần đây một số gia đình đón Tết khác trước như thế nào?",
  ["A. Không ăn gì","B. Đi du lịch trong dịp Tết thay vì chỉ ở nhà","C. Không gặp ai","D. Làm việc nhiều hơn"],1),
 ("Câu 18. Hầu hết người Việt vẫn xem Tết là dịp gì?",
  ["A. Dịp để làm việc","B. Dịp để mua sắm","C. Dịp để sum họp và nghỉ ngơi","D. Dịp để học thi"],2)]),

("【7 会社・人物】 Xe điện（電気自動車・VinFast）",
"VinFast là một công ty của Việt Nam, sản xuất xe ô tô và xe máy chạy bằng điện. Khác với xe thường, xe điện không dùng xăng mà dùng điện để chạy. Vì không đốt xăng nên xe điện không thải khói, ít gây ô nhiễm và tốt cho môi trường. Xe điện cũng chạy êm và ít gây tiếng ồn hơn xe thường. Trong những năm gần đây, ngày càng nhiều người Việt mua xe điện để đi lại trong thành phố. Không chỉ có ô tô và xe máy, một số thành phố còn dùng xe buýt điện để chở khách. Để giúp người dùng, công ty và thành phố đã làm thêm nhiều nơi sạc điện cho xe ở các khu nhà, siêu thị và bãi đỗ xe. Nhà nước cũng có một số chính sách giúp người mua xe điện. Không chỉ bán trong nước, VinFast còn đem xe đi bán ở nước ngoài, trong đó có cả thị trường khó tính như Mỹ và châu Âu. Đây là một điều đáng tự hào với nhiều người Việt Nam. Tuy nhiên, xe điện vẫn còn một số khó khăn. Ví dụ như giá còn hơi cao, nơi sạc điện chưa có ở khắp nơi, và sạc đầy pin cũng cần thời gian. Một số người còn lo về việc thay pin sau nhiều năm sử dụng. Dù vậy, nhiều người tin rằng cùng với thời gian, các khó khăn này sẽ được giải quyết, và xe điện sẽ ngày càng phổ biến hơn trong tương lai.",
"ビンファストはベトナムの会社で、電気で走る自動車やバイクを生産している。普通の車と違い、電気自動車はガソリンではなく電気を使って走る。ガソリンを燃やさないので、電気自動車は煙を出さず、汚染が少なく、環境に良い。電気自動車は静かに走り、普通の車より騒音も少ない。ここ数年、街なかの移動のために電気自動車を買うベトナム人がますます増えている。自動車やバイクだけでなく、一部の都市は客を運ぶのに電気バスも使っている。利用者を助けるため、会社や都市は、住宅地・スーパー・駐車場に車の充電場所を多く作った。国も電気自動車を買う人を助けるいくつかの政策を出している。国内で売るだけでなく、ビンファストは外国にも車を売っていて、その中にはアメリカやヨーロッパのような厳しい市場もある。これは多くのベトナム人にとって誇らしいことだ。ただし、電気自動車にはまだいくつかの難しさがある。たとえば値段がまだやや高いこと、充電場所がどこにでもあるわけではないこと、満充電にも時間がかかることだ。長年使ったあとの電池の交換を心配する人もいる。それでも、時とともにこれらの難しさは解決され、電気自動車は将来ますます広まると、多くの人が信じている。",
[("Câu 19. Xe điện chạy bằng gì?",
  ["A. Bằng xăng","B. Bằng điện","C. Bằng nước","D. Bằng gió"],1),
 ("Câu 20. Vì sao xe điện tốt cho môi trường?",
  ["A. Vì chạy rất nhanh","B. Vì rất rẻ","C. Vì không thải khói, ít gây ô nhiễm","D. Vì rất to"],2),
 ("Câu 21. Khó khăn của xe điện hiện nay là gì?",
  ["A. Giá còn cao và nơi sạc điện chưa có khắp nơi","B. Không ai muốn mua","C. Không bán ở nước ngoài","D. Thải nhiều khói"],0)]),

("【8 会社・人物】 Khởi nghiệp（若者の起業）",
"Ngày nay, nhiều bạn trẻ Việt Nam mơ ước tự mở công ty của riêng mình, gọi là khởi nghiệp. Thay vì chỉ đi làm thuê, họ muốn tự làm chủ và thực hiện ý tưởng của mình. Nhiều người bắt đầu từ những việc nhỏ, ví dụ như bán hàng trên mạng, làm ứng dụng điện thoại, mở quán cà phê hay bán đồ ăn. Có bạn trẻ làm ra sản phẩm mới, cũng có bạn nghĩ ra cách phục vụ tiện hơn cho khách. Để khởi nghiệp thành công, bạn trẻ cần có ý tưởng tốt, chịu khó học hỏi và không sợ thất bại. Họ cũng cần biết quản lý tiền và hiểu khách hàng muốn gì. Một số trường đại học và công ty lớn cũng giúp đỡ bằng cách dạy kiến thức, cho mượn tiền và giới thiệu người có kinh nghiệm để hướng dẫn. Nhờ internet, bạn trẻ ngày nay cũng dễ học hỏi và bán hàng ra nhiều nơi hơn trước. Tuy nhiên, khởi nghiệp không hề dễ. Nhiều công ty mới phải đóng cửa chỉ sau một, hai năm vì thiếu tiền, thiếu kinh nghiệm hoặc vì sản phẩm chưa hợp với khách. Dù vậy, những người từng thất bại nói rằng họ học được rất nhiều và sẽ thử lại với cách làm tốt hơn. Tinh thần dám nghĩ dám làm của người trẻ được nhiều người đánh giá cao, vì nó giúp tạo ra việc làm mới và giúp xã hội phát triển.",
"今、多くのベトナムの若者は、自分自身の会社を立ち上げること、いわゆる起業を夢見ている。ただ雇われて働くのではなく、自分が主になって自分のアイデアを実現したいと考えている。多くの人は小さなことから始める。たとえばネットでの販売、電話のアプリ作り、喫茶店や食べ物の販売などだ。新しい製品を作る若者もいれば、客にもっと便利なサービスを考え出す若者もいる。起業を成功させるには、若者は良いアイデアを持ち、苦労して学び、失敗を恐れないことが必要だ。お金を管理し、客が何を望むかを理解することも必要だ。一部の大学や大企業も、知識を教えたり、お金を貸したり、指導するために経験者を紹介したりして助けている。インターネットのおかげで、今の若者は以前より学びやすく、多くの場所に売りやすい。ただし、起業は決して簡単ではない。多くの新しい会社が、お金や経験の不足、または製品が客に合わないために、わずか一、二年で店を閉じる。それでも、失敗を経験した人は、とても多くを学んだので、よりよいやり方でまた挑戦すると言う。若者の挑戦する精神は、新しい仕事を生み社会の発展を助けるので、多くの人に高く評価されている。",
[("Câu 22. Khởi nghiệp nghĩa là gì?",
  ["A. Đi làm thuê cho công ty lớn","B. Tự mở công ty của riêng mình","C. Đi du học","D. Nghỉ việc ở nhà"],1),
 ("Câu 23. Để khởi nghiệp thành công cần gì?",
  ["A. Ý tưởng tốt, chịu khó học hỏi, không sợ thất bại","B. Thật nhiều tiền lẻ","C. Không cần học","D. Sợ thất bại"],0),
 ("Câu 24. Vì sao nhiều công ty mới phải đóng cửa?",
  ["A. Vì có quá nhiều khách","B. Vì thiếu tiền hoặc thiếu kinh nghiệm","C. Vì nhân viên quá giỏi","D. Vì giá quá thấp"],1)]),

("【9 サービス・客】 Giao đồ ăn（料理の宅配）",
"Bây giờ, đặt đồ ăn qua điện thoại đã trở thành thói quen của nhiều người Việt, nhất là ở thành phố. Người dùng chỉ cần mở ứng dụng, chọn quán và món mình thích, rồi chờ ở nhà hoặc ở công ty. Trên ứng dụng có rất nhiều quán để chọn, từ món Việt đến món nước ngoài. Người dùng cũng có thể xem giá, xem ảnh món ăn và đọc ý kiến của khách trước khi đặt. Sau khi đặt, người giao hàng sẽ mang đồ ăn đến tận nơi, thường vẫn còn nóng. Người dùng có thể trả tiền mặt hoặc trả qua điện thoại. Dịch vụ này rất tiện cho người bận rộn, người ốm hoặc khi trời mưa. Nhờ có dịch vụ giao đồ ăn, nhiều quán nhỏ cũng bán được nhiều hơn vì có thêm khách ở xa. Nhiều người cũng có thêm việc làm nhờ đi giao hàng. Tuy nhiên, dịch vụ này cũng có điểm chưa tốt. Tiền giao hàng đôi khi hơi cao, và vào giờ cao điểm khách có thể phải chờ lâu. Đôi khi đồ ăn mang đến bị nguội hoặc không giống như mong đợi. Ngoài ra, đồ ăn đựng trong nhiều hộp nhựa và túi ni lông cũng gây hại cho môi trường. Vì vậy, một số quán đã bắt đầu dùng hộp giấy thay cho hộp nhựa, và khuyên khách chỉ lấy thìa đũa khi thật sự cần, để cùng nhau bảo vệ môi trường.",
"今、電話で料理を注文することは、特に都市で多くのベトナム人の習慣になった。利用者はアプリを開き、店と好きな料理を選び、家や会社で待つだけでよい。アプリにはベトナム料理から外国料理まで、選べる店がとても多い。利用者は注文する前に、値段を見たり、料理の写真を見たり、客の意見を読んだりできる。注文したあと、配達の人が料理をその場所まで、たいていまだ温かいうちに持ってきてくれる。利用者は現金でも電話でも支払える。このサービスは、忙しい人、病気の人、雨のときにとても便利だ。料理宅配のおかげで、多くの小さな店も、遠くの客が増えて以前より多く売れている。配達をすることで仕事が増える人も多い。ただし、このサービスにも良くない点がある。配達料が時々少し高く、混む時間帯には客が長く待つこともある。料理が冷めて届いたり、思っていたものと違ったりすることもある。さらに、料理が多くのプラスチック容器やビニール袋に入っていることも環境に害を与える。だから、一部の店は環境を一緒に守るために、プラスチック容器の代わりに紙の容器を使い始め、本当に必要なときだけスプーンや箸をもらうよう客に勧めている。",
[("Câu 25. Người dùng đặt đồ ăn như thế nào?",
  ["A. Tự đến quán nấu","B. Mở ứng dụng, chọn quán và món, rồi chờ","C. Gửi thư cho quán","D. Gọi taxi"],1),
 ("Câu 26. Dịch vụ giao đồ ăn tiện cho ai?",
  ["A. Người bận rộn, người ốm, khi trời mưa","B. Chỉ người giàu","C. Chỉ trẻ em","D. Chỉ người ở quê"],0),
 ("Câu 27. Một số quán làm gì để bảo vệ môi trường?",
  ["A. Tăng tiền giao hàng","B. Dùng hộp giấy thay cho hộp nhựa","C. Đóng cửa","D. Bán đắt hơn"],1)]),

("【10 サービス・客】 Học trực tuyến（オンライン学習）",
"Trong những năm gần đây, học trực tuyến ngày càng phổ biến ở Việt Nam. Người học chỉ cần một chiếc máy tính hoặc điện thoại và mạng internet là có thể học ở bất cứ đâu. Có nhiều lớp học trên mạng, từ tiếng Anh, tiếng Nhật cho đến nấu ăn, vẽ hay lập trình. Người học có thể học trực tiếp với thầy cô qua màn hình, hoặc xem lại video bài giảng vào lúc rảnh. Học trực tuyến giúp người học tiết kiệm thời gian đi lại và có thể học ngay tại nhà. Học phí trên mạng đôi khi cũng rẻ hơn lớp học bình thường, và có nhiều bài học miễn phí. Vì vậy, nhiều người đi làm bận rộn vẫn có thể vừa làm vừa học thêm. Học xong một số khóa học, người học còn được cấp giấy chứng nhận. Tuy nhiên, học trực tuyến cũng cần người học phải tự giác. Nếu không chăm chỉ, người học dễ bỏ giữa chừng vì không có ai nhắc nhở. Ngoài ra, học một mình qua màn hình đôi khi khó hỏi bài, ít gặp bạn bè và dễ mỏi mắt nếu ngồi quá lâu. Có lúc bài học cũng bị gián đoạn vì mạng yếu. Vì thế, nhiều người chọn cách vừa học trên mạng vừa học ở lớp để vừa tiện vừa hiệu quả. Dù còn vài khó khăn, học trực tuyến đã mở ra nhiều cơ hội học tập mới cho mọi người.",
"ここ数年、ベトナムではオンライン学習がますます一般的になっている。学ぶ人は、パソコンか電話とインターネットさえあれば、どこでも学べる。英語や日本語から、料理、絵、プログラミングまで、ネット上にはたくさんの講座がある。学ぶ人は画面ごしに先生から直接学ぶことも、暇なときに授業の動画を見直すこともできる。オンライン学習は、通う時間を節約でき、家でそのまま学べる。ネットの授業料は普通の教室より安いこともあり、無料の授業も多い。そのため、忙しく働く多くの人も、働きながら学び足すことができる。いくつかの講座を学び終えると、修了証ももらえる。ただし、オンライン学習は学ぶ人の自覚も必要だ。まじめにやらないと、注意してくれる人がいないので途中でやめてしまいやすい。さらに、画面ごしに一人で学ぶと、時に質問しにくく、友達にもあまり会えず、長く座ると目も疲れやすい。通信が弱くて授業が途切れることもある。だから、便利さと効果の両方のために、ネットでも学び教室でも学ぶ、という方法を選ぶ人も多い。まだいくつかの難しさはあるが、オンライン学習はみんなに新しい学びの機会をたくさん開いた。",
[("Câu 28. Để học trực tuyến cần có gì?",
  ["A. Máy tính hoặc điện thoại và mạng internet","B. Chỉ cần sách","C. Phải đến trường","D. Cần nhiều tiền mặt"],0),
 ("Câu 29. Học trực tuyến có lợi gì?",
  ["A. Phải đi lại nhiều","B. Tiết kiệm thời gian đi lại, học lúc rảnh","C. Luôn đắt hơn","D. Không ai dạy"],1),
 ("Câu 30. Điểm khó của học trực tuyến là gì?",
  ["A. Quá nhiều bạn bè","B. Lớp học quá gần","C. Cần tự giác, dễ bỏ giữa chừng nếu không chăm chỉ","D. Không có lớp để học"],2)]),
]

doc=Document()
st=doc.styles['Normal']; st.font.name=FONT; st.element.rPr.rFonts.set(qn('w:eastAsia'),FONT)

para(doc,"実用ベトナム語検定 4級 読解対策（拡充版）― 時事の長文10本・30問",size=14.5,bold=True,
     color=(0x1F,0x49,0x7D),align=WD_ALIGN_PARAGRAPH.CENTER,after=4)
para(doc,"長文10本×各3問＝30問　／　練習用（実在の特定記事ではなく、最近の傾向をもとにした文）",
     size=9.5,color=(0x66,0x66,0x66),align=WD_ALIGN_PARAGRAPH.CENTER,after=10)

head(doc,"Ⅰ. Bài đọc (やさしいベトナム語・長文10本)")
for title,vn,jp,qs in P:
    para(doc,title,size=11.5,bold=True,after=2)
    para(doc,vn,size=11,after=10)

doc.add_paragraph()
head(doc,"Ⅱ. Câu hỏi đọc hiểu (読解問題・全30問)")
para(doc,"上の本文を読み、正しい答えを1つ選びなさい。",size=10,color=(0x66,0x66,0x66),after=8)
answers=[]
for title,vn,jp,qs in P:
    para(doc,title.split("】")[0]+"】",size=10.5,bold=True,color=(0x55,0x55,0x55),after=2)
    for q,opts,ai in qs:
        para(doc,q,size=11,bold=True,after=2)
        for o in opts: para(doc,"   "+o,size=11,after=1)
        num=q.split(".")[0].replace("Câu ","")
        answers.append((num,opts[ai].split(".")[0]))
        doc.add_paragraph()
ans_line="【解答】 "+"  /  ".join("%s:%s"%(n,a) for n,a in answers)
para(doc,ans_line,size=11,bold=True,color=(0xC0,0,0),after=6)

doc.add_page_break()
head(doc,"Ⅲ. 日本語訳（長文10本）")
for title,vn,jp,qs in P:
    para(doc,title,size=11.5,bold=True,after=2)
    para(doc,jp,size=11,after=10)

doc.add_paragraph()
head(doc,"Ⅳ. 4級の単語・文法")
para(doc,"▼ 単語",size=12,bold=True,color=(0x1F,0x49,0x7D),after=4)
vocab=[
("thanh toán","支払う"),("tiền mặt","現金"),("quét mã QR","QRコードを読み取る"),
("tiền lẻ","小銭"),("uy tín","信用"),("chất lượng","品質"),
("nhân viên","従業員"),("tiết kiệm","節約する"),("tập trung","集中する"),
("đồng nghiệp","同僚"),("cô đơn","さびしい"),("kết hợp","組み合わせる"),
("sức khỏe","健康"),("đều đặn","規則正しく"),("tinh thần","気持ち・精神"),
("bị thương","けがをする"),("khách du lịch","観光客"),("phố cổ","旧市街"),
("đồ lưu niệm","おみやげ"),("phát triển","発展する"),("môi trường","環境"),
("sum họp","集まる・団らんする"),("kỷ niệm","思い出"),("sản xuất","生産する"),
("ô nhiễm","汚染"),("sạc điện","充電する"),("khởi nghiệp","起業する"),
("làm chủ","主になる・経営する"),("thất bại","失敗する"),("kinh nghiệm","経験"),
("giao hàng","配達する"),("giờ cao điểm","ピーク時"),("hộp nhựa","プラスチック容器"),
("học trực tuyến","オンライン学習"),("tự giác","自覚的に行う"),("hiệu quả","効果的"),
]
tbl=doc.add_table(rows=1,cols=2); tbl.style='Table Grid'
for c,t in zip(tbl.rows[0].cells,["ベトナム語","意味"]):
    setf(c.paragraphs[0].add_run(t),size=10.5,bold=True)
for vi,jp_ in vocab:
    cells=tbl.add_row().cells
    setf(cells[0].paragraphs[0].add_run(vi),size=11)
    setf(cells[1].paragraphs[0].add_run(jp_),size=11)

doc.add_paragraph()
para(doc,"▼ 文法（やさしい言い方）",size=12,bold=True,color=(0x1F,0x49,0x7D),after=4)
gram=[
("ngày càng + 形容詞","「ますます〜」。例: ngày càng nhiều người（ますます多くの人）"),
("chỉ cần + 動詞","「〜するだけでよい」。例: chỉ cần mở ứng dụng（アプリを開くだけでよい）"),
("vừa A vừa B","「AしながらB／AでもありBでもある」。例: vừa làm vừa học（働きながら学ぶ）"),
("thay vì + 名詞/動詞","「〜の代わりに」。例: thay vì ở nhà（家にいる代わりに）"),
("nhờ (có) + 名詞","「〜のおかげで」。例: nhờ có internet（インターネットのおかげで）"),
("… nên …","「〜なので…」。例: không đốt xăng nên không thải khói"),
("vì vậy / vì thế","「だから・そのため」。前の理由を受けて結果を言う。"),
("tuy nhiên / nhưng","「しかし・ただし」。前と逆のことを続ける。"),
("dù vậy","「それでも」。前を認めつつ違う結論を続ける。"),
("không chỉ A mà còn B","「AだけでなくBも」。例: không chỉ bán trong nước mà còn bán ở nước ngoài"),
]
for pat,exp in gram:
    para(doc,"・"+pat,size=11,bold=True,after=1)
    para(doc,"    "+exp,size=10.5,after=6)

out=r"C:\Users\baru1\Desktop\ベトナム語\ベトナム語検定\4級\4級読解対策_時事10本30問.docx"
doc.save(out)
print("SAVED:",out)
print("本文",len(P),"設問",sum(len(p[3]) for p in P),"単語",len(vocab),"文法",len(gram))
import statistics
lens=[len(p[1]) for p in P]
print("本文の越語文字数 平均",int(statistics.mean(lens)),"最小",min(lens),"最大",max(lens))
print(ans_line)
