# -*- coding: utf-8 -*-
"""4級 重要4分野まとめ（前置詞の相性／後置修飾／固定構文／接続詞）Word生成。
過去問に出た範囲を超えて、4級相当で網羅的に。"""
import os, sys, io
sys.stdout = io.TextIOWrapper(sys.stdout.buffer, encoding='utf-8', errors='replace')
from docx import Document
from docx.shared import Pt, RGBColor
from docx.oxml.ns import qn
from docx.enum.text import WD_ALIGN_PARAGRAPH

OUT = r"C:\Users\baru1\Hub\Cloud\Study\4級対策\4級まとめ_前置詞・後置修飾・構文・接続詞_2026-06-09.docx"
FONT = "Meiryo UI"

doc = Document()

def setfont(run, size=10.5, bold=False, color=None):
    run.font.size = Pt(size); run.font.bold = bold
    run.font.name = FONT
    rpr = run._element.get_or_add_rPr()
    rf = rpr.find(qn('w:rFonts'))
    if rf is None:
        rf = rpr.makeelement(qn('w:rFonts'), {}); rpr.append(rf)
    for a in ('w:ascii', 'w:hAnsi', 'w:eastAsia', 'w:cs'):
        rf.set(qn(a), FONT)
    if color: run.font.color.rgb = color

NAVY = RGBColor(0x1F, 0x4E, 0x79)
PINK = RGBColor(0xC2, 0x18, 0x5B)

def h1(text):
    p = doc.add_paragraph(); r = p.add_run(text); setfont(r, 16, True, NAVY)
    p.space_after = Pt(4)
def h2(text):
    p = doc.add_paragraph(); r = p.add_run(text); setfont(r, 13, True, RGBColor(0x15,0x65,0xC0))
    p.paragraph_format.space_before = Pt(8); p.paragraph_format.space_after = Pt(2)
def para(text, size=10.5, color=None):
    p = doc.add_paragraph(); r = p.add_run(text); setfont(r, size, False, color)
    return p

def table(headers, rows, widths=None):
    t = doc.add_table(rows=1, cols=len(headers)); t.style = "Table Grid"
    for i, hd in enumerate(headers):
        c = t.rows[0].cells[i]; c.text = ""
        rr = c.paragraphs[0].add_run(hd); setfont(rr, 10, True, RGBColor(0xFF,0xFF,0xFF))
        shd = c._tc.get_or_add_tcPr().makeelement(qn('w:shd'), {qn('w:fill'): "1565C0", qn('w:val'): "clear"})
        c._tc.get_or_add_tcPr().append(shd)
    for row in rows:
        cells = t.add_row().cells
        for i, val in enumerate(row):
            cells[i].text = ""
            for j, seg in enumerate(str(val).split("||")):
                p = cells[i].paragraphs[0] if j == 0 else cells[i].add_paragraph()
                rr = p.add_run(seg)
                # 1列目（ベトナム語/構文）はピンク強調
                setfont(rr, 10, (i == 0), PINK if i == 0 else None)
    return t

# ===== 表紙 =====
h1("ベトナム語検定 4級 重要4分野 まとめ")
para("前置詞の相性／後置修飾／固定構文／接続詞の分類　― 4級相当で網羅的に", 10.5, RGBColor(0x60,0x60,0x60))
para("※過去問に出た語だけでなく、4級でよく使う形を広く収録。最終確認はHuong先生に。", 9.5, RGBColor(0x90,0x90,0x90))

# ===== 1. 前置詞の相性 =====
h1("Ⅰ. 前置詞の相性（動詞・形容詞＋前置詞）")
para("ベトナム語は「動詞・形容詞」ごとに組む前置詞が決まっている。誤り問題で最頻出。", 10)

h2("① với（〜と・〜に｜相手・対象）")
table(["ベトナム語", "意味", "例文"], [
 ["kết hôn với", "〜と結婚する", "Cô ấy kết hôn với một bác sĩ."],
 ["so sánh với", "〜と比べる", "Đừng so sánh mình với người khác."],
 ["giống với", "〜と似ている", "Em ấy giống với mẹ."],
 ["khác với", "〜と違う", "Khác với tôi, anh ấy thích cà phê."],
 ["quen với", "〜に慣れる／知り合い", "Tôi đã quen với công việc mới."],
 ["hài lòng với", "〜に満足する", "Tôi hài lòng với kết quả."],
 ["thất vọng với", "〜に失望する", "Tôi thất vọng với cách làm đó."],
 ["kiên nhẫn với", "〜に忍耐強い", "Hãy kiên nhẫn với trẻ con."],
 ["đối diện với", "〜の向かいに", "Nhà tôi đối diện với bệnh viện."],
 ["giao tiếp với", "〜とやり取りする", "Học tiếng Anh để giao tiếp với khách."],
 ["liên lạc với", "〜と連絡を取る", "Tôi sẽ liên lạc với anh sau."],
 ["hợp tác với", "〜と協力する", "Công ty hợp tác với đối tác Nhật."],
 ["cạnh tranh với", "〜と競争する", "Các hãng cạnh tranh với nhau."],
 ["đồng ý với", "〜に同意する", "Tôi đồng ý với ý kiến của bạn."],
])

h2("② về（〜について｜話題・方向）")
table(["ベトナム語", "意味", "例文"], [
 ["nói/viết về", "〜について話す/書く", "Bài báo viết về du lịch Việt Nam."],
 ["nghĩ về", "〜について考える", "Tôi đang nghĩ về tương lai."],
 ["biết / tìm hiểu về", "〜について知る/調べる", "Tôi muốn tìm hiểu về văn hóa Việt."],
 ["lo lắng về", "〜について心配する", "Mẹ lo lắng về sức khỏe của tôi."],
 ["ý kiến về", "〜についての意見", "Bạn có ý kiến gì về việc này?"],
 ["chịu trách nhiệm về", "〜に責任を負う", "Ai chịu trách nhiệm về sai sót này?"],
 ["(hài lòng/thất vọng về)", "満足/失望（về も可）", "Tôi hài lòng về dịch vụ."],
])

h2("③ đến / tới（〜に・〜まで｜到達点）")
table(["ベトナム語", "意味", "例文"], [
 ["quan tâm đến/tới", "〜に関心がある", "Tôi quan tâm đến môi trường."],
 ["chú ý đến", "〜に注意する", "Hãy chú ý đến an toàn."],
 ["ảnh hưởng đến", "〜に影響する", "Thời tiết ảnh hưởng đến mùa màng."],
 ["liên quan đến", "〜に関連する", "Vấn đề này liên quan đến tiền bạc."],
 ["từ … đến …", "〜から〜まで", "Tôi làm từ 9 giờ đến 5 giờ."],
])

h2("④ sang（〜へ｜別の言語・別の側へ）")
table(["ベトナム語", "意味", "例文"], [
 ["dịch sang tiếng …", "〜語に訳す", "Hãy dịch câu này sang tiếng Nhật."],
 ["chuyển sang", "〜に切り替える", "Tôi chuyển sang công việc khác."],
 ["đổi sang", "〜に替える", "Đổi tiền yên sang tiền Việt."],
])

h2("⑤ vào（〜に｜中へ・拠り所）")
table(["ベトナム語", "意味", "例文"], [
 ["cho … vào", "〜を中に入れる", "Cho tài liệu vào tủ."],
 ["tham gia vào", "〜に参加する", "Tôi tham gia vào câu lạc bộ."],
 ["tập trung vào", "〜に集中する", "Hãy tập trung vào bài học."],
 ["dựa vào", "〜に頼る／基づく", "Đừng dựa vào người khác."],
 ["tin vào", "〜を信じる", "Tôi tin vào bản thân."],
 ["phụ thuộc vào", "〜に左右される", "Giá phụ thuộc vào trọng lượng."],
])

h2("⑥ cho（〜に・〜のために｜受け手）")
table(["ベトナム語", "意味", "例文"], [
 ["tặng/gửi cho", "〜に贈る/送る", "Tôi gửi quà cho bạn."],
 ["nói/kể cho", "〜に話す", "Hãy kể cho tôi nghe."],
 ["đưa cho", "〜に渡す", "Đưa cho tôi quyển sách."],
 ["dành cho", "〜のためのもの", "Món quà này dành cho em."],
 ["làm cho", "〜させる／のためにする", "Tin đó làm cho tôi vui."],
])

h2("⑦ để（〜するために｜目的＋動詞）")
table(["ベトナム語", "意味／注意", "例文"], [
 ["để + 動詞", "〜するために", "Tôi học tiếng Việt để làm việc."],
 ["★目的（名詞）は vì / nhằm mục đích", "×để mục đích", "Anh làm vậy nhằm mục đích gì?"],
])

h2("⑧ bằng（〜で｜手段・材料）")
table(["ベトナム語", "意味", "例文"], [
 ["đi bằng", "〜（乗り物）で行く", "Tôi đi bằng xe buýt."],
 ["làm bằng", "〜（材料）で作る", "Cái túi làm bằng da."],
 ["viết bằng", "〜（言語/道具）で書く", "Viết bằng bút chì."],
 ["trả bằng", "〜（手段）で払う", "Trả bằng thẻ được không?"],
])

h2("⑨ theo（〜によると・〜に沿って）")
table(["ベトナム語", "意味", "例文"], [
 ["theo + 名詞 + thì …", "〜によると…", "Theo dự báo thì ngày mai mưa."],
 ["làm theo", "〜に従って行う", "Làm theo hướng dẫn."],
])

# ===== 2. 後置修飾 =====
h1("Ⅱ. 後置修飾（名詞を後ろから説明する）")
para("ベトナム語は説明する語を名詞の【後ろ】に置く（日本語と逆）。読解・語義問題で必須。", 10)
table(["型", "意味", "例（下線部が後置修飾）"], [
 ["名詞 ＋ 動詞句", "〜する／〜である…（関係代名詞を使わない）", "người làm việc trong tổ chức ＝組織で働く人||bác sĩ chữa răng ＝歯を治す医者"],
 ["名詞 ＋ (đã/đang/sẽ) ＋ 動詞", "〜した/している/する…", "thời gian đã qua ＝過ぎ去った時間||người đang đứng kia ＝あそこに立っている人"],
 ["名詞 ＋ mà ＋ 主語＋動詞", "〜（する）…（関係詞 mà）", "quyển sách mà tôi đã mua ＝私が買った本||điều mà mình cho là đúng ＝自分が正しいと思うこと"],
 ["名詞 ＋ 形容詞", "〜な…", "cái áo đẹp ＝きれいな服||căn phòng rộng ＝広い部屋"],
 ["名詞 ＋ của ＋ 名詞", "〜の…", "nhà của tôi ＝私の家||ý kiến của giám đốc ＝社長の意見"],
 ["名詞 ＋ để ＋ 動詞", "〜するための…", "điều để giải thích ＝説明するためのもの||cách để học giỏi ＝上手に学ぶ方法"],
 ["名詞 ＋ nào đó", "ある…・どこかの…（不特定）", "một vùng nào đó ＝ある地域||vì lý do nào đó ＝何かの理由で"],
 ["名詞 ＋ này/đó/kia", "この/その/あの…", "người này ＝この人||việc đó ＝そのこと"],
])

# ===== 3. 固定構文 =====
h1("Ⅲ. 固定構文（覚えれば作れる｜並べ替えの得点源）")
para("4級でよく出る決まった型。意味と語順をセットで暗記する。", 10)
table(["構文", "意味", "例文"], [
 ["càng ngày càng ＋ 形", "日に日に〜", "Trời càng ngày càng lạnh."],
 ["càng A càng B", "AするほどB", "Càng học càng thấy thú vị."],
 ["形/動 ＋ đến nỗi / đến mức ＋ 結果", "〜すぎて…／〜なほど", "Mệt đến nỗi không đứng dậy được."],
 ["(Mặc dù / Tuy / Dù) A nhưng B", "Aだけれど B", "Tuy nhà xa nhưng anh ấy không đi muộn."],
 ["Dù A thì B cũng/vẫn …", "たとえAでもBは…", "Dù khó thì tôi cũng sẽ làm."],
 ["Dù A hay B thì … cũng", "AでもBでも…", "Dù mưa hay nắng thì tôi cũng đi."],
 ["Nếu A thì B", "もしAならB", "Nếu rảnh thì gọi cho tôi."],
 ["(Do / Vì) A nên B", "AなのでB", "Vì trời mưa nên tôi ở nhà."],
 ["không những A mà còn B (nữa)", "AだけでなくBも", "Cô ấy không những đẹp mà còn thông minh."],
 ["vừa A vừa B", "AしながらB／AでもありB", "Nó vừa ăn vừa xem tivi."],
 ["疑問詞 ＋ cũng（ai/gì/đâu/nào/bao giờ cũng）", "誰でも/何でも/どこでも…", "Ai cũng biết điều đó."],
 ["không ＋ 動 ＋ gì cả", "何も〜ない", "Tôi không hiểu gì cả."],
 ["chưa từng ＋ 動 ＋ bao giờ", "一度も〜したことがない", "Tôi chưa từng đến đó bao giờ."],
 ["trông ＋ 主 ＋ có vẻ ＋ 形", "〜に見える／〜そうだ", "Trông em ấy có vẻ mệt."],
 ["để ＋ 動", "〜するために", "Tôi đến sớm để chuẩn bị."],
 ["trước khi / sau khi ＋ 動", "〜する前に/後に", "Rửa tay trước khi ăn."],
 ["bao nhiêu … bấy nhiêu", "〜だけ…（同じ量）", "Muốn mua bao nhiêu thì bán bấy nhiêu."],
 ["sao … vậy", "〜のとおりに…する", "Nói sao thì làm vậy."],
 ["cứ … là …", "〜すると決まって…", "Cứ mưa to là đường ngập."],
 ["đáng ＋ 動", "〜する価値がある", "Bộ phim này đáng xem."],
 ["khó / dễ ＋ 動", "〜しにくい／しやすい", "Bài này khó hiểu."],
 ["tự ＋ 動 ＋ lấy", "自分で〜する", "Tôi tự làm lấy."],
 ["không được ＋ 動", "〜してはいけない", "Không được hút thuốc ở đây."],
 ["cần / nên / phải ＋ 動", "必要/べき/ねばならない", "Bạn nên nghỉ ngơi."],
 ["có thể nói (rằng) …", "〜と言える", "Có thể nói đây là biểu tượng của Nhật."],
 ["thế nào … cũng", "どうあっても〜", "Thế nào tôi cũng sẽ đến."],
])

# ===== 4. 接続詞の分類 =====
h1("Ⅳ. 接続詞の分類（読解の空所で得点）")
para("前後が「順接・逆接・並列・時系列…」のどれかを見て選ぶ。", 10)
table(["種類", "接続詞", "意味・例"], [
 ["順接（だから）", "nên / cho nên / vì vậy / vì thế / do đó", "Trời mưa nên tôi ở nhà.（だから）"],
 ["逆接（しかし）", "nhưng / tuy nhiên / thế nhưng / song", "Tôi mệt, tuy nhiên vẫn làm.（しかし）"],
 ["理由（なぜなら）", "vì / bởi vì / do / tại vì", "Tôi nghỉ vì bị ốm.（なので）"],
 ["並列・追加（その上/また/同時に）", "và / hơn nữa / ngoài ra / bên cạnh đó / đồng thời", "Món này ngon, hơn nữa lại rẻ.（その上）"],
 ["時系列（最初に→次→最後）", "đầu tiên / trước hết → sau đó / tiếp theo / rồi → cuối cùng", "Đầu tiên…, sau đó…, cuối cùng…"],
 ["例示（例えば/特に）", "ví dụ như / chẳng hạn (như) / đặc biệt là / cụ thể là", "ví dụ như Hồ Tây, Hồ Hoàn Kiếm（例えば）"],
 ["言い換え・結論（つまり/要するに）", "tức là / nói cách khác / nói chung / tóm lại", "Tóm lại, chúng ta nên tiết kiệm.（要するに）"],
 ["対比（一方で/逆に）", "mặt khác / ngược lại / trong khi đó", "Anh thích trà; ngược lại, tôi thích cà phê.（逆に）"],
])

para("", 8)
para("【使い方】まず①〜④を声に出して読み、例文を5回ずつ。問3（前置詞）・問5（構文）・問6（接続詞）は、ここを覚えるだけで得点が安定します。", 10, RGBColor(0x2E,0x7D,0x32))

# Normalスタイルにもフォント設定
st = doc.styles["Normal"]
st.font.name = FONT; st.font.size = Pt(10.5)
st.element.rPr.rFonts.set(qn('w:eastAsia'), FONT)

doc.save(OUT)
print("saved:", OUT)
print("paragraphs:", len(doc.paragraphs), "tables:", len(doc.tables))
