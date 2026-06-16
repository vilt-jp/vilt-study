# -*- coding: utf-8 -*-
"""Huong先生 並び替え35問の「文法詳細解説＋単語(品詞・漢越語つき)」をWord化。
品詞・全例文の日本語訳・先生の手書き追記(原本PDF p1-5)を反映した詳細版。Goodnote取込用。
元データ: huong-lesson-2026-05-02/data.js REARRANGE_QUESTIONS(35問のans)
手書き原本: Desktop/ベトナム語/ベトナム語検定/4級/Đề luyện thi cấp 4 ( ngày 26 tháng 4).pdf
フォントMeiryo UI(latin/ea/cs)で越日両対応。"""
from docx import Document
from docx.shared import Pt, RGBColor
from docx.oxml.ns import qn
from docx.enum.text import WD_ALIGN_PARAGRAPH

FONT = "Meiryo UI"
BLUE = (0x1F, 0x49, 0x7D)
GRAY = (0x66, 0x66, 0x66)
RED = (0xB0, 0x2A, 0x1F)      # 先生の手書き
GREEN = (0x1B, 0x5E, 0x20)    # 品詞

def setf(run, size=None, bold=None, color=None, italic=None):
    run.font.name = FONT
    rPr = run._element.get_or_add_rPr()
    rF = rPr.find(qn('w:rFonts'))
    if rF is None:
        rF = rPr.makeelement(qn('w:rFonts'), {}); rPr.insert(0, rF)
    for a in ('w:ascii', 'w:hAnsi', 'w:eastAsia', 'w:cs'): rF.set(qn(a), FONT)
    if size is not None: run.font.size = Pt(size)
    if bold is not None: run.font.bold = bold
    if italic is not None: run.font.italic = italic
    if color is not None: run.font.color.rgb = RGBColor(*color)

def para(doc, text="", size=11, bold=False, color=None, align=None, after=5, before=0, italic=None, indent=None):
    p = doc.add_paragraph()
    if align is not None: p.alignment = align
    p.paragraph_format.space_after = Pt(after); p.paragraph_format.space_before = Pt(before)
    if indent is not None: p.paragraph_format.left_indent = Pt(indent)
    r = p.add_run(text); setf(r, size=size, bold=bold, color=color, italic=italic)
    return p

def head(doc, text, size=14):
    return para(doc, text, size=size, bold=True, color=BLUE, after=8, before=10)

doc = Document()
st = doc.styles['Normal']; st.font.name = FONT
st.element.rPr.rFonts.set(qn('w:eastAsia'), FONT)

para(doc, "Huong先生 並び替え35問 ― 文法詳細解説（品詞・例文・日本語訳つき）", size=15, bold=True,
     color=BLUE, align=WD_ALIGN_PARAGRAPH.CENTER, after=3)
para(doc, "4級対策・ベトナム語学習用／先生の手書きメモを反映", size=9.5, color=GRAY,
     align=WD_ALIGN_PARAGRAPH.CENTER, after=4)
para(doc, "凡例：【品詞】＝その表現の文法上の種類　／　★手書き＝先生が原本に手書きで足したメモ", size=9, color=GRAY,
     align=WD_ALIGN_PARAGRAPH.CENTER, after=10)

head(doc, "Ⅰ. 文法パターン")

# 各項目: (番号タイトル, 品詞, 説明, [(越文, 和訳), ...], [手書きメモ, ...])
grammar = [
 ("1. 〜だけれども（譲歩）：Dù / Mặc dù / Tuy … (nhưng) … vẫn / cũng",
  "接続詞（譲歩）。Dù/Mặc dù/Tuy＝接続詞、vẫn/cũng＝副詞（それでも・〜も）",
  "逆のことが起きても結果は変わらない、という言い方。後ろに vẫn（それでも）や cũng（〜も）が来る。Tuy が一番かたい言い方。",
  [("Dù gia đình có phản đối thì tôi vẫn cưới cô ấy.", "家族が反対しても、私はそれでも彼女と結婚する。"),
   ("Dù bố mẹ không đồng ý nhưng Lan vẫn đi du học.", "両親が同意しなくても、ランはそれでも留学する。"),
   ("Mặc dù thời tiết xấu nhưng trận đấu vẫn được diễn ra.", "天気が悪かったけれども、試合はそれでも行われた。"),
   ("Tuy nhà xa nhưng Hiền chưa bao giờ đi làm muộn.", "家が遠いけれども、ヒエンは一度も遅刻したことがない。"),
   ("Dù có nghèo thì tôi cũng không bao giờ ăn trộm.", "たとえ貧しくても、私は決して盗みはしない。")],
  ["★手書き：型は「Dù＋S₁＋có＋V₁＋thì＋S₂＋(vẫn)…」。có は「たとえ〜しても」を強める（có phản đối＝反対しても）。",
   "★手書き：「Dù … nhưng …」の形もある。"]),

 ("2. 〜であろうとなかろうと：Dù … hay (không) … thì … cũng",
  "接続詞（譲歩・二者選択）。dù…hay…＝〜であれ〜であれ、cũng＝副詞（どちらにせよ）",
  "二つのどちらでも結果は同じ、という言い方。",
  [("Khi được tặng quà, dù thích hay không thì chúng ta nên nói lời cảm ơn.", "贈り物をもらったら、好きでも嫌いでも、私たちはお礼を言うべきだ。"),
   ("Dù bận rộn hay rảnh rỗi thì tôi cũng gọi điện cho bố mẹ hằng ngày.", "忙しくても暇でも、私は毎日両親に電話する。")],
  []),

 ("3. 〜によって…される：do（行為者）＋動詞　／　動詞＋ra",
  "do＝前置詞（受け身の動作主＝英語の by）。ra＝方向補語（補助的な語）",
  "英語の受け身に近い。「誰がやったか」を do の後ろに置く。動詞＋ra は「外に生み出す」イメージ。",
  [("Mì ăn liền do Ando Momofuku phát minh ra.", "即席麺は安藤百福によって発明された。"),
   ("Bản báo cáo này do chị Vy viết.", "この報告書はヴィーさんによって書かれた。")],
  ["★手書き：do＝by（〜によって）。",
   "★手書き：ra は「外へ生み出す」イメージ。phát minh ra（発明する）／ nghĩ ra（思いつく）／ sáng tạo ra（創り出す）。"]),

 ("4. 〜なので…（理由→結果）：(Do / Vì) … nên …",
  "接続詞（因果）。Do/Vì＝〜なので（理由）、nên＝だから（結果）",
  "前に理由、nên のあとに結果を置く。",
  [("Anh ấy thường xuyên bỏ bữa sáng nên bị đau dạ dày.", "彼は頻繁に朝食を抜くので、胃が痛くなる。"),
   ("Do lười học nên tôi phải thi lại môn tiếng Pháp.", "勉強を怠けたので、私はフランス語の科目を再試験しなければならない。")],
  []),

 ("5. 〜だけでなく…も：không những … mà còn … (nữa)",
  "接続詞（並列・添加）。không những＝〜だけでなく、mà còn＝その上〜も、nữa＝さらに（副詞）",
  "二つの良い点・特徴を並べる。後ろが名詞のときは là（＝である）を2回入れる。",
  [("Anh Kimura không những nấu được món Nhật mà còn (nấu được) món ăn Trung Quốc nữa.", "キムラさんは日本料理が作れるだけでなく、中華料理も作れる。"),
   ("Trời không những mưa mà còn lạnh nữa.", "空は雨が降るだけでなく、寒くもある。"),
   ("Cô ấy không những là một diễn viên mà còn là một nhà văn.", "彼女は俳優であるだけでなく、作家でもある。")],
  ["★手書き：名詞をつなぐときは「không những là N₁ mà còn là N₂」と là が2回必要。",
   "★手書き：動詞のときは V₁ … mà còn (V₂)。例の (nấu được) は省略されることもある。"]),

 ("6. 例えば〜：ví dụ như / … như …",
  "連語（例示）。ví dụ＝名詞（例）＋ như＝前置詞（〜のような）",
  "例を挙げるときの言い方。như だけでも「〜のような」と例を示せる。",
  [("Hà Nội có nhiều hồ nước tự nhiên ví dụ như Hồ Tây, Hồ Hoàn Kiếm.", "ハノイには天然の湖が多く、例えばタイ湖やホアンキエム湖などがある。")],
  ["★手書き：「ví dụ A, B, C …」と並べる。",
   "★手書き：như でも例示できる。例）Anh thích ăn món Nhật như sushi, …（寿司のような日本料理が好き）。"]),

 ("7. どんな〜でも：bất cứ + 名詞 + nào cũng ／ bất cứ điều gì ／ bất cứ thứ gì ／ bất cứ lúc nào",
  "連語（全称）。bất cứ＝どんな〜でも、nào＝疑問詞、cũng＝副詞（みな）",
  "「どれでも・何でも・いつでも」と全部を指す。",
  [("Tại Việt Nam, bất cứ căn bếp nào cũng có nước mắm.", "ベトナムでは、どの台所にも必ず魚醤（ヌクマム）がある。"),
   ("Trước khi làm bất cứ điều gì, bạn phải suy nghĩ.", "何かをする前に、あなたはよく考えなければならない。"),
   ("Giờ đây, chúng ta có thể mua sắm trực tuyến bất cứ thứ gì vào bất cứ lúc nào.", "今や、私たちはいつでも何でもオンラインで買い物ができる。")],
  ["★手書き：型は「bất cứ＋N＋nào＋cũng …」。",
   "★手書き：mua sắm trực tuyến＝オンライン（online）で買い物。",
   "★手書き：suy nghĩ về việc＋V／ suy nghĩ về＋N（〜について考える）。"]),

 ("8. 〜次第・〜に応じて：tùy theo / tùy thuộc vào　（成句：tùy cơ ứng biến）",
  "tùy theo＝前置詞句（〜に従って）、tùy thuộc vào＝動詞句（〜に依る）",
  "条件によって結果が変わる、という言い方。",
  [("Tùy theo trọng lượng và kích thước, giá tiền thay đổi.", "重さと大きさに応じて、値段は変わる。"),
   ("Câu trả lời tùy thuộc vào cách suy nghĩ của mỗi người.", "答えは人それぞれの考え方次第だ。"),
   ("Tùy cơ ứng biến.", "臨機応変に（その場に応じて対応する）。")],
  ["★手書き：Tùy cơ ứng biến は四字熟語（臨機応変）。"]),

 ("9. 〜するほど・〜するあまり：đến mức / đến nỗi",
  "連語（程度）。đến＝〜に至る＋mức/nỗi＝程度（名詞）",
  "程度が大きいことを強調する。",
  [("Bài tập về nhà nhiều đến mức học sinh muốn khóc.", "宿題が多すぎて、生徒は泣きたくなるほどだ。"),
   ("Em trai tôi mải chơi game đến nỗi quên ăn.", "弟はゲームに夢中になりすぎて、食べるのを忘れるほどだ。"),
   ("Câu chuyện của Lan làm tôi cảm động đến mức rơi nước mắt.", "ランの話は、私を涙が出るほど感動させた。")],
  []),

 ("10. 〜するために：để + 動詞　（＋ trước khi＝〜する前に）",
  "để＝接続詞・前置詞（目的）。trước khi＝接続詞（〜の前に）",
  "目的を表す。trước khi は「〜する前に」。",
  [("Hành khách phải xếp hàng chờ hàng giờ tại sân bay để làm thủ tục.", "乗客は手続きのために、空港で何時間も並んで待たなければならない。")],
  []),

 ("11. 〜だけ…する（同じ量）：bao nhiêu … bấy nhiêu",
  "呼応表現（数量）。bao nhiêu＝どれだけ、bấy nhiêu＝それだけ（指示詞）",
  "前の量と同じだけ後ろもする。",
  [("Chị muốn mua bao nhiêu thì tôi sẽ bán cho chị bấy nhiêu.", "あなたが買いたいだけ、私はあなたに売ってあげます。")],
  []),

 ("12. 〜のとおりに…する：sao … vậy",
  "呼応表現。sao＝どのように、vậy＝そのように（指示詞）",
  "言われたとおりに行う。",
  [("Giám đốc yêu cầu sao thì tôi làm vậy.", "社長が求めるとおりに、私はやります。")],
  []),

 ("13. 〜すると決まって…：cứ … là …",
  "接続詞（条件・反復）。cứ＝〜しさえすれば、là＝すると必ず",
  "あることが起きると必ず次が起きる。",
  [("Cứ mưa to là đường phố Hà Nội bị ngập.", "大雨が降ると決まって、ハノイの街は水浸しになる。")],
  []),

 ("14. 〜させる（引き起こす）：khiến / làm / làm cho",
  "使役動詞。khiến / làm / làm cho＝〜を…させる",
  "原因が結果を引き起こす言い方。",
  [("Tiếng còi xe tải khiến tôi giật mình.", "トラックのクラクションの音が、私をびっくりさせた。"),
   ("Câu chuyện của Lan làm tôi cảm động.", "ランの話は、私を感動させた。"),
   ("Covid-19 đã làm cho ngành du lịch tê liệt hoàn toàn.", "コロナは観光業を完全に麻痺させた。")],
  []),

 ("15. もし〜なら…（すべき）：nếu … thì (nên) …",
  "接続詞（条件）。nếu＝もし、thì＝ならば、nên＝〜すべき（助動詞）",
  "条件と、それに対する助言・結果。",
  [("Nếu anh muốn giải quyết một cách nhanh chóng thì nên thuê luật sư.", "もし早く解決したいなら、弁護士を雇うべきだ。")],
  []),

 ("16. 自分で〜する：tự … lấy",
  "tự＝副詞（自分で）、lấy＝補助詞（自分の力で、を強める）",
  "自分の力で行うことを強める。",
  [("Tôi tự may quần áo lấy.", "私は自分で服を縫う。")],
  []),

 ("17. 〜と言える：có thể nói",
  "連語（前置き）。có thể＝〜できる（助動詞）＋ nói＝言う（動詞）",
  "「〜と言ってよい」という前置き。文頭に来る。",
  [("Có thể nói núi Phú Sĩ là biểu tượng của Nhật Bản.", "富士山は日本の象徴だと言える。"),
   ("Có thể nói Covid-19 đã thay đổi thói quen người tiêu dùng.", "コロナは消費者の習慣を変えたと言える。")],
  ["★手書き：có thể nói は文頭に来る。"]),

 ("18. その他のポイント：chưa bao giờ / không bao giờ",
  "副詞（頻度）。chưa bao giờ＝一度も〜ない、không bao giờ＝決して〜ない",
  "経験・習慣の否定を表す。",
  [("Tuy nhà xa nhưng Hiền chưa bao giờ đi làm muộn.", "家が遠いけれども、ヒエンは一度も遅刻したことがない。"),
   ("Dù có nghèo thì tôi cũng không bao giờ ăn trộm.", "たとえ貧しくても、私は決して盗みはしない。")],
  []),
]

for title, pos, desc, exs, notes in grammar:
    para(doc, title, size=11.5, bold=True, after=2, before=6)
    para(doc, "【品詞】" + pos, size=10, color=GREEN, after=2, indent=6)
    if desc:
        para(doc, desc, size=10.5, after=2, indent=6)
    for vi, jp in exs:
        p = doc.add_paragraph(); p.paragraph_format.space_after = Pt(1); p.paragraph_format.left_indent = Pt(10)
        r1 = p.add_run("・" + vi + "\n"); setf(r1, size=10.5)
        r2 = p.add_run("　" + jp); setf(r2, size=10, color=GRAY)
    for nt in notes:
        para(doc, nt, size=10, color=RED, after=1, indent=10)

# ===== 手書きノート（原本 p4-5）=====
doc.add_page_break()
head(doc, "Ⅱ. 先生の手書きノート（原本 4〜5ページ）")
para(doc, "原本の後半は、先生が手書きで足した文法ノートです。以下に全て書き起こしました。", size=10, color=GRAY, after=8)

hnotes = [
 ("① bất cứ ＋ 名詞 ＋ nào ＋ cũng（どんな〜でも）",
  [("Bất cứ người nào cũng phải nộp thuế.", "どんな人でも税金を払わなければならない。"),
   ("Bất cứ món nào cũng ngon.", "どの料理でも美味しい。"),
   ("Bất cứ người nào cũng không được vào.", "どんな人も入ってはいけない。")]),
 ("② cho ＋ N₁ ＋ vào ＋ N₂（N₁ を N₂ に入れる）",
  [("Trứng thì không được cho vào lò vi sóng.", "卵は電子レンジに入れてはいけない。"),
   ("Em không được cho trứng vào lò vi sóng.", "卵（N₁）を電子レンジ（N₂）に入れてはいけない。")]),
 ("③ lấy ＋ N₁ ＋ từ trong ＋ N₂ ＋ ra ngoài（N₂ の中から N₁ を取り出す）",
  [("Lấy bia từ trong tủ lạnh ra.", "冷蔵庫の中からビールを取り出す。"),
   ("Lấy tài liệu từ trong ngăn kéo ra.", "引き出しの中から書類を取り出す。")]),
 ("④ cách の使い方",
  [("cách ＋ 動詞 ＝ 〜のやり方（例：cách suy nghĩ ＝ 考え方）", ""),
   ("Từ Hà Nội đến HCM cách bao xa?", "ハノイからホーチミンまでどれくらい（の距離）離れていますか？")]),
]
for title, items in hnotes:
    para(doc, title, size=11, bold=True, color=RED, after=2, before=5)
    for vi, jp in items:
        p = doc.add_paragraph(); p.paragraph_format.space_after = Pt(1); p.paragraph_format.left_indent = Pt(10)
        r1 = p.add_run("・" + vi + ("\n" if jp else "")); setf(r1, size=10.5)
        if jp:
            r2 = p.add_run("　" + jp); setf(r2, size=10, color=GRAY)

para(doc, "⑤ các と những の違い（手書きの図より）", size=11, bold=True, color=RED, after=2, before=6)
para(doc, "・các ＝ あるグループの「全体」を指す。　例）các em ＝ 生徒みんな（全員）。", size=10.5, after=1, indent=10)
para(doc, "・những ＝ その中の「特定の一部」を指す。　例）những em nhóm A ＝ （みんなの中の）グループAの生徒たち。", size=10.5, after=1, indent=10)
para(doc, "（図：大きい円 các em の中に nhóm A・nhóm B・nhóm C の小さい円。những em nhóm A は A の集団だけを指す。）", size=9.5, color=GRAY, after=4, indent=10)

para(doc, "⑥ その他の手書きメモ（原本1ページ・動詞穴埋め問題の欄）", size=11, bold=True, color=RED, after=2, before=6)
para(doc, "・cùng ＋ N ／ với ＋ N ＝ どちらも「〜と一緒に」。後ろに名詞が来る。", size=10.5, after=4, indent=10)

# ===== 単語 =====
doc.add_page_break()
head(doc, "Ⅲ. 単語（品詞・漢越語つき）")
para(doc, "「品詞」列＝名・動・形・副など。「漢越語」列＝漢字に対応するベトナム語。"
         "— は固有語（漢越語でない）。（diễn演）のように一部だけが漢越語の場合はその部分を示す。",
     size=9.5, color=GRAY, after=8)

# (越, 意味, 品詞, 漢越語)
vocab = [
 ("phản đối", "反対する", "動", "反対"),
 ("cưới", "結婚する（娶る）", "動", "—"),
 ("cố gắng hết sức", "全力を尽くす", "動", "—"),
 ("thay đổi", "変える・変わる", "動", "—"),
 ("tình hình", "情勢・状況", "名", "情形"),
 ("phát minh", "発明する", "動", "発明"),
 ("bản báo cáo", "報告書", "名", "報告（báo cáo）／本（bản）"),
 ("diễn viên", "俳優", "名", "演員"),
 ("nhà văn", "作家", "名", "文（văn）"),
 ("hồ (nước)", "湖", "名", "湖（hồ）"),
 ("căn bếp", "台所", "名", "—"),
 ("nước mắm", "魚醤（ヌクマム）", "名", "—"),
 ("suy nghĩ", "考える", "動", "推（suy）"),
 ("trọng lượng", "重さ", "名", "重量"),
 ("kích thước", "大きさ・寸法", "名", "—"),
 ("giá tiền", "値段", "名", "価（giá）"),
 ("câu trả lời", "答え", "名", "—"),
 ("cách suy nghĩ", "考え方", "名", "格（cách）・推（suy）"),
 ("bài tập về nhà", "宿題", "名", "習（tập）"),
 ("khóc", "泣く", "動", "—"),
 ("quên", "忘れる", "動", "—"),
 ("hành khách", "乗客", "名", "行客"),
 ("xếp hàng", "列に並ぶ", "動", "—"),
 ("sân bay", "空港", "名", "—"),
 ("làm thủ tục", "手続きをする", "動", "手続（thủ tục）"),
 ("giám đốc", "社長・部長", "名", "監督"),
 ("yêu cầu", "要求する", "動", "要求"),
 ("mưa to", "大雨", "名", "—"),
 ("đường phố", "通り", "名", "—"),
 ("bị ngập", "水浸しになる", "動", "—"),
 ("giật mình", "びっくりする", "動", "—"),
 ("cảm động", "感動する", "動", "感動"),
 ("rơi nước mắt", "涙を流す", "動", "—"),
 ("ngành du lịch", "観光業", "名", "遊歴（du lịch）"),
 ("tê liệt", "麻痺する", "動", "—"),
 ("hoàn toàn", "完全に", "副", "完全"),
 ("giải quyết", "解決する", "動", "解決"),
 ("thuê", "雇う・借りる", "動", "—"),
 ("luật sư", "弁護士", "名", "律師"),
 ("đồng ý", "同意する", "動", "同意"),
 ("đi du học", "留学する", "動", "遊学（du học）"),
 ("thời tiết", "天気", "名", "時節"),
 ("trận đấu", "試合", "名", "—"),
 ("diễn ra", "行われる", "動", "演（diễn）"),
 ("nghèo", "貧しい", "形", "—"),
 ("ăn trộm", "盗む", "動", "—"),
 ("tặng quà", "贈り物をする", "動", "贈（tặng）"),
 ("bận rộn", "忙しい", "形", "—"),
 ("rảnh rỗi", "暇", "形", "—"),
 ("gọi điện", "電話する", "動", "電（điện）"),
 ("hằng ngày", "毎日", "副", "恒（hằng）"),
 ("biểu tượng", "象徴", "名", "表象"),
 ("thói quen", "習慣", "名", "—"),
 ("người tiêu dùng", "消費者", "名", "消（tiêu）"),
 ("thường xuyên", "頻繁に", "副", "常（thường）"),
 ("đau dạ dày", "胃が痛い", "動", "—"),
 ("lười học", "勉強を怠ける", "動", "学（học）"),
 ("thi lại", "追試・再試験", "動", "試（thi）"),
]
tbl = doc.add_table(rows=1, cols=4); tbl.style = 'Table Grid'
for c, t in zip(tbl.rows[0].cells, ["ベトナム語", "意味", "品詞", "漢越語（漢字）"]):
    r = c.paragraphs[0].add_run(t); setf(r, size=10.5, bold=True)
for vi, jp, pos, hv in vocab:
    cells = tbl.add_row().cells
    setf(cells[0].paragraphs[0].add_run(vi), size=11)
    setf(cells[1].paragraphs[0].add_run(jp), size=11)
    setf(cells[2].paragraphs[0].add_run(pos), size=11)
    setf(cells[3].paragraphs[0].add_run(hv), size=11)

out = r"C:\Users\baru1\Desktop\ベトナム語\Huong並び替え_文法詳細解説_品詞例文訳つき.docx"
doc.save(out)
print("SAVED:", out)
print("文法", len(grammar), "手書きノート群", len(hnotes) + 2, "単語", len(vocab))
