# -*- coding: utf-8 -*-
import sys, io, os
sys.stdout = io.TextIOWrapper(sys.stdout.buffer, encoding='utf-8', errors='replace')
from docx import Document
f = r"C:\Users\baru1\Desktop\ベトナム語\Huong並び替え_文法と単語_漢越語つき 追記.docx"
print("exists:", os.path.exists(f))
if os.path.exists(f):
    d = Document(f)
    print("=== PARAGRAPHS ===")
    for p in d.paragraphs:
        t = p.text.rstrip()
        if t:
            print(repr(p.style.name) if p.style else '', "|", t)
    print("=== TABLES:", len(d.tables), "===")
    for ti, tb in enumerate(d.tables):
        print("--- TABLE", ti, "rows", len(tb.rows), "---")
        for r in tb.rows:
            print(" | ".join(c.text.strip().replace(chr(10), " / ") for c in r.cells))
