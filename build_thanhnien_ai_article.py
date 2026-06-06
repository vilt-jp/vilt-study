# -*- coding: utf-8 -*-
"""Thanh Nien記事「Hơn 50% học sinh dùng AI làm bài」学習用Word作成。
構成: ①越語原文 ②読解問題5問 ③(改ページ)日本語訳 ④4級単語・文法解説
フォントは Meiryo UI (latin/ea/cs 全設定で越語・日本語両対応)。"""
from docx import Document
from docx.shared import Pt, RGBColor
from docx.oxml.ns import qn
from docx.enum.text import WD_ALIGN_PARAGRAPH

FONT = "Meiryo UI"

def set_font(run, size=None, bold=None, color=None):
    run.font.name = FONT
    rPr = run._element.get_or_add_rPr()
    rFonts = rPr.find(qn('w:rFonts'))
    if rFonts is None:
        rFonts = rPr.makeelement(qn('w:rFonts'), {})
        rPr.insert(0, rFonts)
    for a in ('w:ascii', 'w:hAnsi', 'w:eastAsia', 'w:cs'):
        rFonts.set(qn(a), FONT)
    if size is not None:
        run.font.size = Pt(size)
    if bold is not None:
        run.font.bold = bold
    if color is not None:
        run.font.color.rgb = RGBColor(*color)

def para(doc, text="", size=11, bold=False, color=None, align=None, space_after=6):
    p = doc.add_paragraph()
    if align is not None:
        p.alignment = align
    p.paragraph_format.space_after = Pt(space_after)
    r = p.add_run(text)
    set_font(r, size=size, bold=bold, color=color)
    return p

def heading(doc, text, size=14, color=(0x1F, 0x49, 0x7D)):
    p = para(doc, text, size=size, bold=True, color=color, space_after=8)
    return p

doc = Document()
# 既定スタイルもMeiryo UI
st = doc.styles['Normal']
st.font.name = FONT
st.element.rPr.rFonts.set(qn('w:eastAsia'), FONT)

# ===== タイトル =====
para(doc, "Hơn 50% học sinh dùng AI làm bài, nói gian lận là điều phổ biến",
     size=16, bold=True, color=(0x1F, 0x49, 0x7D), align=WD_ALIGN_PARAGRAPH.CENTER, space_after=4)
para(doc, "Báo Thanh Niên ・ Tác giả: Ngọc Long ・ Ngày đăng: 28/02/2026 16:19 GMT+7",
     size=9.5, color=(0x66, 0x66, 0x66), align=WD_ALIGN_PARAGRAPH.CENTER, space_after=2)
para(doc, "出典: thanhnien.vn（実用ベトナム語学習用にまとめたもの）",
     size=9, color=(0x66, 0x66, 0x66), align=WD_ALIGN_PARAGRAPH.CENTER, space_after=10)

# ===== ① 越語原文 =====
heading(doc, "Ⅰ. Nguyên văn (ベトナム語原文)")
vn_paras = [
"Đó là kết quả nghiên cứu của Trung tâm nghiên cứu Pew (Mỹ) về xu hướng dùng trí tuệ nhân tạo (AI) của thanh thiếu niên Mỹ, trong bối cảnh AI tác động khắp các ngóc ngách của ngành giáo dục.",
"Cụ thể, 54% học sinh Mỹ từ 13 tới 17 tuổi cho biết đã sử dụng các chatbot AI như ChatGPT, Copilot và Character.ai để làm bài tập trong lớp trong khi 57% dùng các công cụ này để tra cứu thông tin. Hai mục đích khác cũng được đông đảo học sinh Mỹ hướng tới là để giải trí hoặc tìm kiếm niềm vui (47%) và để tóm tắt bài báo, sách vở hoặc video (42%).",
"Xu hướng đáng chú ý là phụ huynh lại khá ủng hộ con cái sử dụng AI, thể hiện qua việc có 58% cha mẹ thấy ổn khi con nhờ AI hỗ trợ làm bài, 23% thấy không ổn và 15% không chắc về vấn đề này. Tuy nhiên nếu con cái muốn tìm tới AI để trò chuyện hay tìm lời khuyên, tỷ lệ ủng hộ lần lượt giảm xuống chỉ còn 28% và 18%.",
"Kết quả khảo sát 1.458 cặp học sinh và phụ huynh ở Mỹ (mỗi cặp gồm 1 học sinh và 1 phụ huynh) từ ngày 25.9 đến 9.10.2025 cho biết thêm, mức độ ứng dụng AI trong giới trẻ có sự khác biệt đáng kể. Trong khi 44% em cho biết chỉ dùng AI cho “một vài” hay “số ít” bài tập, thì 10% thú nhận nhờ chatbot hỗ trợ làm toàn bộ hoặc hầu hết bài vở được giao.",
"“Chúng tôi chắc chắn việc sử dụng chatbot AI để hỗ trợ làm bài tập đang trở thành một thói quen phổ biến với thanh thiếu niên”, bà Colleen McClain, nghiên cứu viên cấp cao ở Trung tâm nghiên cứu Pew và là đồng tác giả bài nghiên cứu, nhận định.",
"Nghiên cứu về một chủ đề nào đó là lý do được gọi tên nhiều nhất khi học sinh Mỹ tìm đến chatbot AI, với 48% tỷ lệ đồng tình. Theo sau là nhờ giải toán (43%) và chỉnh sửa văn bản (35%). Đáng chú ý, 26% học sinh được khảo sát chia sẻ rằng các chatbot AI “cực kỳ” hoặc “rất” hữu ích, 25% nói thấy “phần nào” hữu ích.",
"Chỉ 3% học sinh nhận định phần phản hồi từ AI không hề hoặc không quá hữu ích, và 45% còn lại không dùng chatbot AI hỗ trợ làm bài nên không nêu ý kiến.",
"Ngoài ra, có tới 34% thanh thiếu niên đồng tình rằng AI sẽ dạy kỹ năng cho các em tốt hơn là một con người thực thụ, trong khi 26% cho rằng nó sẽ dạy kém hơn con người.",
"Một điểm đáng chú ý là 59% giới trẻ Mỹ cho rằng việc dùng AI để gian lận là điều diễn ra thường xuyên ở trường các em. Ngược lại, 14% thông tin học sinh trường các em hiếm khi hoặc chưa từng gian lận bằng chatbot AI, trong khi số còn lại nói rằng không chắc liệu điều này có xảy ra trong môi trường học đường hay không.",
"Điều đặc biệt là nếu chỉ tính riêng nhóm học sinh có sử dụng AI để hỗ trợ làm bài, có đến 76% em tin rằng các bạn đồng trang lứa chung trường đang dùng chatbot AI để gian lận - ít nhất là ở tần suất “thỉnh thoảng”.",
"“Mọi người đang lệ thuộc quá mức vào AI để làm bài tập hay hỏi những câu căn bản”, một nam học sinh chia sẻ trong nghiên cứu.",
"Theo tờ The New York Times, các phát hiện nêu trên được công bố trong bối cảnh toàn nước Mỹ đang diễn ra nhiều cuộc tranh luận gay gắt về sự lan rộng của những AI tạo sinh - công nghệ cho phép người dùng tạo văn bản, hình ảnh chân thực nhất có thể.",
"Trong khi nhóm ủng hộ cho rằng trường cần dạy học sinh cách sử dụng, đánh giá chatbot AI để giúp các em thích ứng với nhu cầu của thị trường lao động, nhóm đối lập lại cảnh báo rằng chatbot AI có thể tạo ra thông tin sai lệch, làm suy giảm tư duy phản biện, tiềm ẩn nguy cơ tác động tiêu cực tới sức khỏe tinh thần và tạo điều kiện cho gian lận nở rộ trong môi trường học đường.",
"Chẳng hạn, kết quả nghiên cứu khả năng đọc hiểu được Hội đồng khảo thí và nhà xuất bản ĐH Cambridge phối hợp với Microsoft Research chỉ ra rằng, những học sinh được yêu cầu ghi chú lại sau khi đọc văn bản mà không được đụng tới chatbot AI có kết quả đọc hiểu tốt hơn so với nhóm dùng chatbot để hỗ trợ hiểu nội dung văn bản.",
"Trước thực tế này, trả lời đài CBS News, nhà tâm lý học Joshua Goodman, phó giáo sư ĐH Southern Oregon (Mỹ), khuyến cáo phụ huynh cần chú ý các dấu hiệu cảnh báo, chẳng hạn như việc thanh thiếu niên liên tục dùng AI, để cho công nghệ thay thế tư duy phản biện của mình hoặc khi các em có biểu hiện trầm cảm.",
]
for t in vn_paras:
    para(doc, t, size=11, space_after=6)

# ===== ② 読解問題 =====
doc.add_paragraph()
heading(doc, "Ⅱ. Câu hỏi đọc hiểu (読解問題・5問)")
para(doc, "Đọc bài trên và chọn đáp án đúng nhất. (本文を読み、最も適切な答えを1つ選びなさい)",
     size=10, color=(0x66, 0x66, 0x66), space_after=8)

questions = [
 ("Câu 1. Theo bài báo, bao nhiêu phần trăm học sinh Mỹ từ 13 đến 17 tuổi đã dùng chatbot AI để làm bài tập?",
  ["A. 42%", "B. 47%", "C. 54%", "D. 59%"]),
 ("Câu 2. Bao nhiêu phần trăm phụ huynh thấy ổn khi con nhờ AI hỗ trợ làm bài?",
  ["A. 23%", "B. 28%", "C. 58%", "D. 76%"]),
 ("Câu 3. Bao nhiêu phần trăm giới trẻ cho rằng việc dùng AI để gian lận diễn ra thường xuyên ở trường các em?",
  ["A. 14%", "B. 34%", "C. 59%", "D. 76%"]),
 ("Câu 4. Nghiên cứu của ĐH Cambridge phối hợp với Microsoft Research cho thấy điều gì?",
  ["A. Học sinh dùng chatbot AI có kết quả đọc hiểu tốt hơn",
   "B. Học sinh ghi chú mà không dùng chatbot AI có kết quả đọc hiểu tốt hơn",
   "C. Chatbot AI không ảnh hưởng đến khả năng đọc hiểu",
   "D. Phụ huynh phản đối việc dùng AI ở trường"]),
 ("Câu 5. Ai là tác giả của bài báo này?",
  ["A. Colleen McClain", "B. Joshua Goodman", "C. Ngọc Long", "D. The New York Times"]),
]
for q, opts in questions:
    para(doc, q, size=11, bold=True, space_after=2)
    for o in opts:
        para(doc, "   " + o, size=11, space_after=1)
    doc.add_paragraph()

para(doc, "【解答】 Câu 1: C  /  Câu 2: C  /  Câu 3: C  /  Câu 4: B  /  Câu 5: C",
     size=11, bold=True, color=(0xC0, 0x00, 0x00), space_after=6)

# ===== 改ページ → ③ 日本語訳 =====
doc.add_page_break()
heading(doc, "Ⅲ. 日本語訳（全文）")
para(doc, "見出し: 「中高生の半数以上がAIで宿題、不正は当たり前だと言う」", size=11, bold=True, space_after=8)
jp_paras = [
"これは、AIが教育のあらゆる場面に影響を及ぼすなか、アメリカの若者の人工知能（AI）利用傾向についてアメリカのピュー研究所が行った調査の結果である。",
"具体的には、13歳から17歳のアメリカの生徒の54%が、ChatGPT、Copilot、Character.aiといったAIチャットボットを授業の課題をやるために使ったことがあると答え、57%は情報を調べるためにこれらの道具を使っている。多くの生徒が向かうもう2つの目的は、娯楽・楽しみを求めるため（47%）と、記事・本・動画を要約するため（42%）である。",
"注目すべき傾向は、保護者がむしろ子どものAI利用にかなり賛成していることだ。58%の親が子どもがAIに宿題を手伝ってもらうことを問題ないと考え、23%はよくないと考え、15%はこの件についてよく分からないとした。ただし、子どもがおしゃべりや助言を求めるためにAIを使う場合、賛成の割合はそれぞれ28%、18%にまで下がる。",
"アメリカの生徒と保護者1,458組（各組は生徒1人と保護者1人）を2025年9月25日から10月9日まで調査した結果はさらに、若者のAI活用の度合いにはかなりの差があることを示している。44%の子は「いくつか」または「ごく少数」の課題にだけAIを使うと答えた一方、10%は出された課題の全部またはほとんどをチャットボットに手伝ってもらったと打ち明けた。",
"「AIチャットボットを課題の補助に使うことが、若者にとって一般的な習慣になりつつあるのは確かです」と、ピュー研究所の上級研究員で本研究の共同著者であるColleen McClain氏は述べた。",
"あるテーマについて調べることが、アメリカの生徒がAIチャットボットを使う理由として最も多く挙げられ、48%が同意した。続いて数学を解いてもらうこと（43%）、文章を直すこと（35%）であった。注目すべきは、調査対象の生徒の26%がAIチャットボットは「非常に」または「とても」役に立つと答え、25%が「ある程度」役に立つと答えたことだ。",
"AIからの返答が全く、またはあまり役に立たないとした生徒はわずか3%で、残りの45%はAIチャットボットを課題の補助に使っていないため意見を述べていない。",
"さらに、34%もの若者が、AIは本物の人間よりも上手に技能を教えてくれるだろうと同意した一方、26%は人間より下手に教えるだろうと考えている。",
"注目すべき点として、アメリカの若者の59%が、AIを使って不正をすることは自分の学校でよく起きていると考えている。逆に、14%は自分の学校の生徒がAIチャットボットで不正をすることはめったにない、または一度もないと伝え、残りはこれが学校という場で起きているかどうか分からないと答えた。",
"特筆すべきは、AIを課題の補助に使っている生徒だけに絞ると、76%もの子が、同じ学校の同年代の仲間が—少なくとも「ときどき」の頻度で—AIチャットボットを使って不正をしていると信じていることだ。",
"「みんな課題をやるためや基本的なことを質問するために、AIに頼りすぎています」と、ある男子生徒が調査の中で語った。",
"ニューヨーク・タイムズ紙によれば、上記の発見は、生成AI—利用者ができる限り本物らしい文章や画像を作れる技術—の広がりについて、アメリカ全土で激しい議論が起きているなかで公表された。",
"賛成派が、労働市場のニーズに適応できるよう学校は生徒にAIチャットボットの使い方・評価の仕方を教えるべきだと考える一方、反対派は、AIチャットボットは誤った情報を生み出し、批判的思考を弱め、精神的健康に悪影響を及ぼす恐れがあり、学校という場で不正がはびこる土壌を作りかねないと警告している。",
"たとえば、ケンブリッジ大学の試験委員会・出版部がMicrosoft Researchと共同で行った読解力の研究は、文章を読んだ後にAIチャットボットに触れずにメモを取るよう求められた生徒の方が、内容理解の補助にチャットボットを使った生徒よりも読解の成績がよかったことを示している。",
"こうした実態を前に、CBSニュースの取材に答えたサザンオレゴン大学（アメリカ）准教授で心理学者のJoshua Goodman氏は、たとえば若者が絶え間なくAIを使い、自分の批判的思考を技術に置き換えてしまう、あるいは抑うつの兆候を見せるといった警告のサインに保護者は注意する必要があると勧告している。",
]
for t in jp_paras:
    para(doc, t, size=11, space_after=6)

# ===== ④ 4級単語・文法 =====
doc.add_paragraph()
heading(doc, "Ⅳ. 4級相当の単語・文法解説")
para(doc, "▼ 単語（記事に出てきた重要語）", size=12, bold=True, color=(0x1F,0x49,0x7D), space_after=4)
vocab = [
 ("nghiên cứu", "研究／調査（する）"),
 ("xu hướng", "傾向・トレンド"),
 ("phụ huynh", "保護者（＝cha mẹ）"),
 ("học sinh", "生徒"),
 ("gian lận", "不正・カンニング（をする）"),
 ("phổ biến", "一般的な・広く普及した"),
 ("công cụ", "道具・ツール"),
 ("tra cứu", "（情報を）調べる・検索する"),
 ("giải trí", "娯楽・気晴らし"),
 ("tóm tắt", "要約（する）"),
 ("ủng hộ", "支持する・賛成する"),
 ("khảo sát", "調査（する）"),
 ("thói quen", "習慣"),
 ("hữu ích", "役に立つ・有益な"),
 ("đáng chú ý", "注目すべき"),
 ("lệ thuộc", "依存する・頼りきる"),
 ("cảnh báo", "警告（する）"),
 ("khuyến cáo", "勧告する・忠告する"),
 ("thường xuyên", "頻繁に・しょっちゅう"),
 ("kết quả", "結果"),
]
tbl = doc.add_table(rows=1, cols=2)
tbl.style = 'Table Grid'
hc = tbl.rows[0].cells
for c, txt in zip(hc, ["ベトナム語", "意味（日本語）"]):
    c.text = ""
    r = c.paragraphs[0].add_run(txt); set_font(r, size=10.5, bold=True)
for vn, jp in vocab:
    cells = tbl.add_row().cells
    r0 = cells[0].paragraphs[0].add_run(vn); set_font(r0, size=11)
    r1 = cells[1].paragraphs[0].add_run(jp); set_font(r1, size=11)

doc.add_paragraph()
para(doc, "▼ 文法（4級でよく出る表現）", size=12, bold=True, color=(0x1F,0x49,0x7D), space_after=4)
grammar = [
 ("Trong khi A ... (thì) B ...", "「Aである一方、Bは…」と対比を表す。例: 54% dùng AI làm bài trong khi 57% dùng để tra cứu.（54%が宿題に使う一方、57%は検索に使う）"),
 ("Tuy nhiên ...", "「しかし／ただし」。前の文と逆の内容を続ける。例: Tuy nhiên, tỷ lệ ủng hộ giảm xuống.（ただし賛成の割合は下がる）"),
 ("Ngược lại ...", "「逆に」。反対の事実を述べる。例: Ngược lại, 14% nói hiếm khi gian lận.（逆に14%はめったに不正しないと言う）"),
 ("cho rằng / cho biết", "cho rằng=「～だと考える」、cho biết=「～だと述べる・知らせる」。例: 59% cho rằng gian lận là phổ biến.（59%が不正は普通だと考える）"),
 ("Chẳng hạn (như) ...", "「たとえば」。例を挙げる。例: Chẳng hạn, nghiên cứu của Cambridge...（たとえばケンブリッジの研究は…）"),
 ("lần lượt", "「それぞれ・順に」。複数の数字を順番に対応させる。例: giảm xuống còn 28% và 18%.（それぞれ28%と18%に下がる）"),
 ("ít nhất", "「少なくとも」。例: ít nhất là ở tần suất “thỉnh thoảng”（少なくとも「ときどき」の頻度で）"),
 ("đáng + 動詞/形容詞", "「～に値する・～すべき」。例: đáng chú ý（注目すべき）, đáng tin（信頼できる）"),
]
for pat, exp in grammar:
    para(doc, "・" + pat, size=11, bold=True, space_after=1)
    para(doc, "    " + exp, size=10.5, space_after=6)

out = r"C:\Users\baru1\Desktop\ベトナム語\Thanh Nien_AIで宿題_学習まとめ.docx"
doc.save(out)
print("SAVED:", out)
print("段落(越語):", len(vn_paras), " 設問:", len(questions), " 訳段落:", len(jp_paras), " 単語:", len(vocab), " 文法:", len(grammar))
